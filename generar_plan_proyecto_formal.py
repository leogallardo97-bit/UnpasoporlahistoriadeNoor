# -*- coding: utf-8 -*-
"""
Generador de Documento: Plan de Proyecto - Noor x Al-Andalus (Hito 1)
Consultoria Estrategica y Direccion de Proyectos PAT-INV
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# Directorio de salida solicitado
OUTPUT_DIR = r"C:\Users\leoga\Desktop\Noor"

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB5)
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
            run.font.size = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def create_wbs_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Fase', 'Actividad Clave', 'Entregable']
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1F4E78')
        hdr[i]._tc.get_or_add_tcPr().append(shading)

    data = [
        ("1. Extraccion", "Auditoria de G:\Mi unidad\\Noor_ 2026_archivos", "Inventario de Fuentes Digitales"),
        ("2. Procesamiento", "Ingesta en NotebookLM y Analisis Historico", "Base de Conocimiento Sintetizada"),
        ("3. Propuesta", "Generacion de Modelos de Investigacion y Menu", "Plan Detallado por Siglo / Recetario"),
        ("4. Control", "Revision por expertos y validacion academica", "Informe de Calidad y Rigor")
    ]
    for fase, act, ent in data:
        row = table.add_row().cells
        row[0].text = fase
        row[1].text = act
        row[2].text = ent
    doc.add_paragraph()

def build_cover(doc, title, subtitle):
    # Espaciado inicial
    for _ in range(5): doc.add_paragraph()
    
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.size = Pt(16)
        run.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for _ in range(10): doc.add_paragraph()
    
    p4 = doc.add_paragraph("Director de Proyecto: Antigravity AI\nConsultoria Estrategica en Patrimonio e Innovacion\n2026")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p4.runs:
        run.font.size = Pt(10)
    doc.add_page_break()

def generate_project_plan():
    doc = Document()
    
    # Portada Profesional
    build_cover(doc, "PLAN DE PROYECTO: NOOR X AL-ANDALUS", "Hito 1: Auditoria Historica e Integracion de Archivo Digital")

    # 1. RESUMEN EJECUTIVO
    set_heading(doc, "1. Resumen Ejecutivo", 1)
    add_paragraph(doc, (
        "El presente proyecto constituye la piedra angular de la estrategia 2026 de Noor. "
        "Este hito inicial busca tender un puente tecnologico y metodologico entre el vasto "
        "legado documental de Al-Andalus y el proceso creativo y operativo de la empresa. "
        "Mediante el procesamiento sistematico del repositorio central (G:\\) y el uso de "
        "AI de vanguardia, transformaremos datos historicos en activos de innovacion gastronomica "
        "y academica de alto impacto."
    ))

    # 2. OBJETIVOS ESTRATEGICOS
    set_heading(doc, "2. Objetivos Estrategicos", 1)
    add_bullet(doc, "Implementar una auditoria exhaustiva del repositorio G:\\Mi unidad\\Noor_ 2026_archivos.")
    add_bullet(doc, "Contextualizar cronologicamente los hitos de Al-Andalus para su aplicacion en la narrativa Noor.")
    add_bullet(doc, "Optimizar el flujo de datos mediante NotebookLM para generar respuestas analiticas precisas.")
    add_bullet(doc, "Asegurar la trazabilidad bibliografica para garantizar el rigor academico del Proyecto Noor.")

    # 3. ESTRUCTURA DE TRABAJO (WBS)
    set_heading(doc, "3. Estructura de Trabajo (WBS)", 1)
    add_paragraph(doc, "La ejecucion se divide en fases secuenciales que garantizan la integridad de la informacion:")
    create_wbs_table(doc)

    # 4. RECURSOS
    set_heading(doc, "4. Recursos y Herramientas", 1)
    add_paragraph(doc, "El exito del proyecto se apoya en una infraestructura hibrida:", bold=True)
    add_bullet(doc, "Repositorio Central: Unidad G:\\ (Informes semanales, archivos digitales y analisis bibliometricos).")
    add_bullet(doc, "Motores de IA: Antigravity AI (Direccion), Claude (Redaccion) y NotebookLM (Analisis documental).")
    add_bullet(doc, "Infraestructura Local: Directorio de trabajo en C:\\Users\\leoga\\Desktop\\Noor.")

    # 5. CRONOGRAMA Y CONTROL DE CALIDAD
    set_heading(doc, "5. Cronograma y Control de Calidad", 1)
    add_paragraph(doc, (
        "El cronograma propuesto para este hito es de ejecucion inmediata, con ciclos de revision "
        "bi-semanales. El control de calidad se regira por los estandares de la Qurtuba Academy, "
        "asegurando que cada receta o hito historico propuesto este academicamente referenciado."
    ))
    
    add_paragraph(doc, "Validacion Final:", bold=True)
    add_bullet(doc, "Cruce de datos entre el Archivo Digital Noor y fuentes externas (WOS/Scopus).")
    add_bullet(doc, "Revision de viabilidad tecnica para su implementacion en sala / master oficial.")

    # Salida
    out_path = os.path.join(OUTPUT_DIR, "Plan_Proyecto_Noor_AlAndalus_Hito1.docx")
    doc.save(out_path)
    print(f"Plan de Proyecto generado: {out_path}")

if __name__ == "__main__":
    generate_project_plan()
