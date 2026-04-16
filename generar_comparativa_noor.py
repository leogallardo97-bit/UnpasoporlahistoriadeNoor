# -*- coding: utf-8 -*-
"""
Generador de Documento: Noor x Al-Andalus: Comparativa Temporal y Propuestas de Proyecto
Proyecto Noor
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

OUTPUT_DIR = r"C:\Users\leoga\Desktop\Noor\Carpeta Plan de invest"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
            run.font.size = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

def build_cover(doc, title, subtitle, periodo):
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
    p2 = doc.add_paragraph(periodo)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(12)
    p3 = doc.add_paragraph("Documento Analitico y Propositivo")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.font.size = Pt(11)
        run.italic = True
    p4 = doc.add_paragraph("Proyecto Noor - Elaborado con asistencia de AI - 2026")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


def create_comparativa():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "Noor × Al-Andalus",
        "Comparativa Temporal y Propuestas de Proyecto",
        "Sintesis Estrategica")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "Este documento presenta una comparativa estructurada del enfoque conceptual del Proyecto Noor "
        "con relacion a las diferentes etapas historicas de Al-Andalus, estableciendo conexiones entre "
        "el pasado andalusi y las propuestas de valor presentes y futuras del proyecto. El objetivo es "
        "articular una vision clara de como Noor asimila, transforma y proyecta el legado andalusi "
        "hacia propuestas operativas, narrativas y gastronomicas."
    ))

    set_heading(doc, "1. Comparativa Temporal: Etapas de Al-Andalus y Fases de Noor", 1)
    
    set_heading(doc, "1.1 El Emirato y el Califato Omeya (Siglos VIII - X)", 2)
    add_paragraph(doc, "Contexto Historico:")
    add_bullet(doc, "Consolidacion del poder y llegada de influencias orientales (Ziryab).")
    add_bullet(doc, "Apogeo cultural, arquitectonico y cientifico en Cordoba.")
    add_paragraph(doc, "Integracion en Noor:")
    add_bullet(doc, "Fundacion gastronomica: Rescate de recetas mozarabes y primeras influencias bagdadies.")
    add_bullet(doc, "Propuesta narrativa: La construccion de la identidad, la convivencia inicial y el nacimiento del lujo andalusi.")

    set_heading(doc, "1.2 Las Taifas y los Imperios Norteafricanos (Siglos XI - XIII)", 2)
    add_paragraph(doc, "Contexto Historico:")
    add_bullet(doc, "Fragmentacion politica pero eclosion descentralizada de las artes (reinos de Taifas).")
    add_bullet(doc, "Austeridad e integracion de nuevos sabores y tecnicas traidas por Almoravides y Almohades.")
    add_paragraph(doc, "Integracion en Noor:")
    add_bullet(doc, "Fusion gastronomica: Introduccion de nuevas texturas y especias (citricos, cana de azucar, nuevas variedades de trigo).")
    add_bullet(doc, "Propuesta narrativa: La diversidad regional contemporanea y el refinamiento en unidades mas pequenas (la personalizacion del servicio).")

    set_heading(doc, "1.3 El Reino Nazari y la Supervivencia Morisca (Siglos XIV - XVII)", 2)
    add_paragraph(doc, "Contexto Historico:")
    add_bullet(doc, "La Alhambra como canto del cisne. Refinamiento absoluto previo a la caida.")
    add_bullet(doc, "La literatura aljamiada y la resistencia culinaria morisca despues de 1492.")
    add_paragraph(doc, "Integracion en Noor:")
    add_bullet(doc, "Culminacion gastronomica: La estilizacion final (dulces, hojaldres, mieles, almendras).")
    add_bullet(doc, "Propuesta narrativa: La resistencia de la memoria. Como conservar la cultura a traves del paladar y la experiencia.")

    set_heading(doc, "2. Propuestas de Proyecto (Estrategias Futuras)", 1)

    set_heading(doc, "Propuesta A: Menus Degustacion Tematicos por Periodo", 2)
    add_paragraph(doc, (
        "Desarrollar lineas de menu de investigacion hiper-focalizadas en momentos concretos, permitiendo "
        "comensales viajar en el tiempo:"
    ))
    add_bullet(doc, "Menu 'Ziryab' (El esplendor omeya y las etiquetas de la mesa).")
    add_bullet(doc, "Menu 'Taifas' (Contrastes regionales y sabores magrebies).")
    add_bullet(doc, "Menu 'Albaicin' (La decadencia elegante nazari y el mimetismo morisco).")

    set_heading(doc, "Propuesta B: Laboratorio Documental y Publicaciones (Noor Ediciones)", 2)
    add_paragraph(doc, (
        "Capitalizar la inmensa labor investigadora (como el archivo consolidado) en un sello o formato "
        "de publicacion propia."
    ))
    add_bullet(doc, "Revista/Paper anual con peer-review de historiadores y chefs.")
    add_bullet(doc, "Documental audiovisual enfocado en la 'arqueologia culinaria'.")

    set_heading(doc, "Propuesta C: Integracion Tecnologica y Archivo Abierto", 2)
    add_paragraph(doc, (
        "Utilizar herramientas como NotebookLM (actualmente en curso) para construir una base de conocimiento "
        "interactiva."
    ))
    add_bullet(doc, "Crear una plataforma donde colaboradores puedan explorar los hallazgos en fuentes primarias (ej. manuscritos traducidos).")
    add_bullet(doc, "Uso de visualizaciones de datos (como bibliometria) para mostrar el peso cultural de Cordoba en la academia.")

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El entrelazamiento metodologico que propone Noor trasciende la mera restauracion; se erige como "
        "un instituto de memoria cultural viva. La comparativa temporal asegura que ningun periodo "
        "fundamental quede fuera de la narrativa, mientras que las propuestas de proyecto asientan "
        "las bases para la expansion artistica, academica y comercial de la marca en los proximos anos."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Noor_AlAndalus_Comparativa_y_Propuestas.docx")
    doc.save(out_path)
    print("Guardado exitosamente: " + out_path)
    return out_path


if __name__ == "__main__":
    print("=" * 70)
    print("Generando Documento: Noor x Al-Andalus")
    print("=" * 70)
    path = create_comparativa()
    print("=" * 70)
    print("COMPLETADO. Documento guardado en:")
    print(path)
