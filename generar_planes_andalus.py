
# -*- coding: utf-8 -*-
"""
Generador de Planes de Investigación Al-Ándalus: Siglos XII, XIII y XIV
Proyecto Noor — Investigación Histórica Multidimensional
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"C:\Users\leoga\Desktop\Noor\Carpeta Plan de invest"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS DE FORMATO
# ─────────────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    """Agrega un encabezado con estilo según nivel."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)  # azul oscuro
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
            run.font.size = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_row(table, col1, col2, col3, bold=False, bg_color=None):
    row = table.add_row()
    cells = row.cells
    for i, text in enumerate([col1, col2, col3]):
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if bold:
                    run.bold = True
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            cells[i]._tc.get_or_add_tcPr().append(shading)
    return row

def create_search_table(doc, rows_data):
    """Crea tabla de búsqueda bibliográfica con 3 columnas."""
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    # Encabezado
    hdr = table.rows[0].cells
    headers = ['Bloque', 'Base', 'Cadena de búsqueda']
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A376C')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for bloque, base, cadena in rows_data:
        add_table_row(table, bloque, base, cadena)
    doc.add_paragraph()

def build_cover(doc, title, subtitle, periodo):
    """Portada del documento."""
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)

    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)

    p2 = doc.add_paragraph(periodo)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(12)

    p3 = doc.add_paragraph("Informe histórico-bibliográfico multidimensional")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.font.size = Pt(11)
        run.italic = True

    p4 = doc.add_paragraph("Proyecto Noor — Elaborado con asistencia de Claude (Anthropic) · 2025")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


# ═════════════════════════════════════════════════════════════════════════════
# SIGLO XII
# ═════════════════════════════════════════════════════════════════════════════

def create_siglo_XII():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    build_cover(doc,
        "AL-ÁNDALUS EN EL SIGLO XII",
        "Del dominio Almorávide al Imperio Almohade",
        "c. 1100 – 1199")

    # ── INTRODUCCIÓN ──────────────────────────────────────────────────────────
    set_heading(doc, "Introducción", 1)
    add_paragraph(doc, (
        "El siglo XII representa para Al-Ándalus una era de profundas transformaciones geopolíticas. "
        "La hegemonía almorávide, consolidada en las primeras décadas, comenzó a erosionarse desde mediados "
        "de siglo ante la presión simultánea de dos fuerzas: el avance de los reinos cristianos del norte "
        "y el surgimiento del movimiento almohade en el Magreb. Este período es crucial para comprender "
        "cómo la identidad peninsular islámica se reformuló bajo sucesivas oleadas de poder norteafricano, "
        "y cómo el territorio andalusí fue reduciéndose, aunque con notables pulsiones culturales y científicas."
    ))
    add_paragraph(doc, (
        "El presente informe analiza Al-Ándalus en el siglo XII en cuatro dimensiones —política, económica, "
        "social y cultural—, y proporciona un plan sistematizado de búsqueda bibliográfica en Web of Science "
        "(WOS) y Scopus orientado al Proyecto Noor."
    ))

    # ── 1. POLÍTICO ───────────────────────────────────────────────────────────
    set_heading(doc, "1. Ámbito Político", 1)

    set_heading(doc, "1.1 La consolidación y decadencia almorávide (1100-1145)", 2)
    add_paragraph(doc, (
        "El siglo XII arrancó con los Almorávides (al-Murābiṭūn) como señores indiscutibles de Al-Ándalus. "
        "Su emir Alí ibn Yúsuf (r. 1106-1143) intentó consolidar el control peninsular, pero se enfrentó "
        "a una creciente resistencia interna: los alfaquíes malikíes que inicialmente respaldaron la intervención "
        "comenzaron a desconfiar de un poder que imponía tributos irregulares y no garantizaba la frontera."
    ))
    add_paragraph(doc, (
        "La derrota de Cutanda (1120) frente a Alfonso I el Batallador marcó el inicio del declive militar "
        "almorávide en la Península. La pérdida de Zaragoza en 1118 —que pasó definitivamente a la Corona "
        "de Aragón— fue un golpe territorial irreversible. Las marcas fronterizas se estrecharon y las poblaciones "
        "mozárabes de Andalucía fueron deportadas masivamente al Magreb por sospecha de colaboracionismo."
    ))

    set_heading(doc, "1.2 Las segundas Taifas (1144-1147)", 2)
    add_paragraph(doc, (
        "El colapso almorávide engendró un segundo ciclo de taifas, efímero pero significativo. Entre 1144 y 1147 "
        "surgieron poderes locales en Murcia, Córdoba, Almería y otras ciudades, protagonizados en parte por "
        "líderes sufíes y en parte por notables urbanos. Este período de fragmentación fue incluso más breve que "
        "el del siglo XI, pues la potencia almohade ya se preparaba para cruzar el Estrecho."
    ))

    set_heading(doc, "1.3 La conquista almohade y la nueva unidad (1147-1172)", 2)
    add_paragraph(doc, (
        "Los Almohades (al-Muwaḥḥidūn), movimiento religioso-político fundado por Ibn Túmart en el Atlas marroquí, "
        "cruzaron el Estrecho en 1147 bajo el califa Abd al-Mumin y tomaron Sevilla. En 1172, con la rendición de "
        "los últimos focos de resistencia, toda Al-Ándalus quedó incorporada al gran Imperio Almohade, que se extendía "
        "desde Mauritania hasta la frontera con los reinos cristianos del norte. Córdoba fue reconquistada en 1148. "
        "La nueva capital andalusí fue Sevilla, que viviría bajo los Almohades uno de sus momentos de mayor esplendor."
    ))

    set_heading(doc, "1.4 La presión cristiana: Alfonso VII y sus sucesores", 2)
    add_paragraph(doc, (
        "Mientras Al-Ándalus se reorganizaba bajo el nuevo dominio, los reinos del norte no cesaban su avance. "
        "Alfonso VII de Castilla y León se tituló Emperador de toda España (1135) y saqueó Córdoba en 1146. "
        "Portugal emergió como reino independiente (1143) y comenzó su propia expansión meridional. "
        "La Segunda Cruzada (1147) incluyó como operación periférica la toma de Lisboa por cruzados del norte "
        "de Europa en apoyo al rey Afonso I de Portugal."
    ))

    # ── 2. ECONÓMICO ──────────────────────────────────────────────────────────
    set_heading(doc, "2. Ámbito Económico", 1)

    set_heading(doc, "2.1 La fiscalidad almorávide y almohade", 2)
    add_paragraph(doc, (
        "Los Almorávides mantuvieron la estructura fiscal omeya pero añadieron exacciones para financiar sus "
        "costosas campañas. Los Almohades, en cambio, introdujeron reformas: unificaron las múltiples tasas "
        "locales, reforzaron el control estatal sobre la producción artesanal y el comercio de exportación, "
        "y penalizaron duramente la evasión. A pesar del rigorismo doctrinal, el fisco almohade fue pragmáticamente "
        "racional: necesitaba recursos para sostener un imperio intercontinental."
    ))

    set_heading(doc, "2.2 Comercio y rutas mediterráneas", 2)
    add_paragraph(doc, (
        "El siglo XII fue una época de intensa actividad comercial mediterránea. Almería siguió siendo el gran "
        "puerto exportador de seda, mientras que las repúblicas marineras italianas —Génova, Pisa— establecieron "
        "factorías en el norte de África y copiaron técnicas de navegación y navegación costera de origen andalusí. "
        "El comercio de oro subsahariano, canalizado a través del Magreb hacia Al-Ándalus, dotó a los emires de "
        "una base monetaria sólida. La moneda almohade de oro —dobla— circuló ampliamente en los mercados europeos."
    ))

    set_heading(doc, "2.3 Agricultura y gestión del agua", 2)
    add_paragraph(doc, (
        "La agricultura intensiva de regadío continuó siendo el pilar de la economía rural. Los sistemas de acequias "
        "del Levante, el Valle del Ebro y el Guadalquivir se mantuvieron operativos, pese a las guerras. "
        "Los Almohades promovieron la extensión del cultivo del olivo y el viñedo (para uva de mesa) en zonas "
        "de nueva conquista. La cría de caballos andaluces, demandados por los ejércitos norteafricanos, generó "
        "también ingresos importantes para las élites ganaderas del sur."
    ))

    # ── 3. SOCIAL ─────────────────────────────────────────────────────────────
    set_heading(doc, "3. Ámbito Social", 1)

    set_heading(doc, "3.1 La estructura social bajo los Almohades", 2)
    add_paragraph(doc, (
        "La llegada almohade supuso una ruptura radical con el pluralismo relativo del período de taifas. "
        "El movimiento, de vocación uniformizante y puritana, presionó a las minorías religiosas de forma "
        "sin precedentes en la historia de Al-Ándalus. Judíos y cristianos fueron obligados a elegir entre "
        "la conversión, el exilio o la muerte, especialmente durante los reinados de Abd al-Mumin y Abu Yaqub Yusuf I."
    ))

    set_heading(doc, "3.2 Las minorías: mozárabes y judíos ante la presión almohade", 2)
    add_paragraph(doc, (
        "La comunidad mozárabe de Andalucía, ya debilitada, fue erradicada prácticamente como grupo organizado. "
        "Sus iglesias fueron destruidas o convertidas en mezquitas, y sus élites optaron mayoritariamente "
        "por el exilio hacia los reinos del norte. Las aljamas judías experimentaron una persecución sistemática, "
        "que forzó el éxodo de intelectuales como Maimónides (Córdoba, 1138 – El Cairo, 1204), quien huyó "
        "con su familia primero al norte de África y luego a Oriente. Este éxodo de cerebros fue una pérdida "
        "civilizatoria de primera magnitud para Al-Ándalus."
    ))

    set_heading(doc, "3.3 Las poblaciones rurales y la militarización", 2)
    add_paragraph(doc, (
        "El siglo XII fue un período de intensa militarización social. Las fronteras móviles provocaron "
        "desplazamientos masivos de población campesina, el abandono de aldeas enteras en zonas de conflicto "
        "y la concentración de la población en ciudades amuralladas. El servicio militar en las huestes almohades "
        "se convirtió en un canal de movilidad social para jóvenes de origen bereber y muwallad."
    ))

    # ── 4. CULTURAL ───────────────────────────────────────────────────────────
    set_heading(doc, "4. Ámbito Cultural y Científico", 1)

    set_heading(doc, "4.1 El siglo de los grandes filósofos", 2)
    add_paragraph(doc, (
        "Paradójicamente, el régimen almohade —de vocación antifilosófica en su doctrina más ortodoxa— coincidió "
        "con el florecimiento de las mayores figuras del pensamiento medieval andalusí. Averroes (Ibn Rushd, "
        "Córdoba 1126-1198) desarrolló su extraordinaria obra de comentarios a Aristóteles bajo el patrocinio "
        "del califa Abu Yaqub Yusuf I, que tenía inclinaciones filosóficas. Sus comentarios, traducidos al latín, "
        "transformarían la escolástica europea del siglo XIII."
    ))
    add_paragraph(doc, (
        "Ibn Tufayl (c. 1105-1185), médico y filósofo de Guadix, escribió Hayy ibn Yaqzan, primera novela "
        "filosófica de la literatura universal, que anticipa las controversias sobre el estado de naturaleza. "
        "Maimónides (1138-1204), aunque exiliado, formó su pensamiento en el ambiente intelectual cordobés "
        "antes de la persecución almohade."
    ))

    set_heading(doc, "4.2 La Escuela de Traductores de Toledo", 2)
    add_paragraph(doc, (
        "Un fenómeno crucial del siglo XII, aunque ubicado en territorio ya cristiano, fue el desarrollo de "
        "la Escuela de Traductores de Toledo bajo el arzobispo Raimundo (1125-1151). Aquí, equipos de intelectuales "
        "bilingües —frecuentemente mozárabes, judíos conversos y clérigos europeos— tradujeron sistemáticamente "
        "al latín la ciencia y filosofía árabe: Euclides, Ptolomeo, al-Khwarizmi, Avicena, al-Farabi. "
        "Este proceso de transmisión es uno de los más importantes de la historia intelectual de Occidente."
    ))

    set_heading(doc, "4.3 Arquitectura almohade y arte", 2)
    add_paragraph(doc, (
        "El arte almohade introdujo una estética de sobriedad depurada: los motivos figurativos desaparecieron "
        "y la decoración geométrica y caligráfica alcanzó una complejidad matemática extraordinaria. En Sevilla, "
        "los Almohades construyeron la Gran Mezquita (cuyo alminar, la Giralda, fue iniciado en 1184) y el "
        "Palacio de la Buhaira. El estilo almohade dejó huella duradera en la arquitectura magrebí y en los "
        "edificios hispanos posteriores, influyendo en el arte mudéjar."
    ))

    set_heading(doc, "4.4 Poesía y literatura", 2)
    add_paragraph(doc, (
        "Ibn Quzman (c. 1078-1160) de Córdoba, maestro del zajal, incorporó el árabe dialectal y el romance "
        "a la poesía culta con una modernidad asombrosa. Ibn Bassam de Santarén compiló la Dajira, antología "
        "fundamental de la literatura andalusí. La producción poética no se interrumpió pese a los avatares "
        "políticos, aunque el mecenazgo se desplazó de las cortes de taifas a la corte almohade de Sevilla."
    ))

    # ── 5. BÚSQUEDA BIBLIOGRÁFICA ─────────────────────────────────────────────
    set_heading(doc, "5. Plan de Búsqueda Bibliográfica", 1)
    add_paragraph(doc, (
        "El siguiente plan sistematizado permite recuperar la producción académica especializada en "
        "Web of Science y Scopus, organizado por los cuatro ámbitos analizados más una búsqueda transversal."
    ))
    create_search_table(doc, [
        ("1. Político", "WOS",
         '("Al-Andalus" OR "Islamic Iberia") AND ("Almoravid*" OR "Almohad*" OR "taifa*" OR "Reconquista") AND ("12th century" OR "twelfth century") AND PY=(1980-2025)'),
        ("2. Económico", "WOS/Scopus",
         '("Al-Andalus" OR "Islamic Spain") AND ("trade" OR "taxation" OR "gold" OR "silk" OR "Mediterranean commerce") AND ("Almohad*" OR "12th century")'),
        ("3. Social", "WOS",
         '("Al-Andalus") AND ("Mozarab*" OR "dhimmi" OR "Maimonides" OR "Jewish" OR "persecution" OR "Almohad*") AND ("12th century" OR "twelfth century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Islamic Iberia") AND ("Averroes" OR "Ibn Rushd" OR "Ibn Tufayl" OR "Toledo school" OR "translation movement" OR "Giralda")'),
        ("5. Transversal", "Ambas",
         'TS=("Al-Andalus" OR "Almohad*" OR "Almoravid*" OR "Islamic Iberia") AND PY=(1980-2025) NOT TS=("Ottoman" OR "Crusades" OR "Egypt")'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodológicas", 2)
    items = [
        "Exportar resultados en formato BibTeX o RIS hacia Zotero o Mendeley para deduplicar entre bases de datos.",
        "Revisar directamente: Al-Qantara (CSIC), Journal of Medieval Iberian Studies, Arabica, Medieval Encounters.",
        "Usar el operador de proximidad en WOS: \"Almohad\" NEAR/3 \"economy\" para mayor precisión.",
        "Complementar con Bibliotheca Islamica y el Corpus of Inscriptions of Islamic Iberia (Univ. Granada).",
        "Filtro temporal recomendado: 1980-2025. Idiomas: EN, ES, FR, AR.",
        "Para la Escuela de Toledo: consultar el corpus digital del Proyecto BIUM (Université Paris V).",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    # ── CONCLUSIÓN ────────────────────────────────────────────────────────────
    set_heading(doc, "Conclusión", 1)
    add_paragraph(doc, (
        "El siglo XII en Al-Ándalus fue un período de tensiones constitutivas: entre uniformidad religiosa "
        "y pluralismo intelectual, entre contracción territorial y expansión cultural. La paradoja almohade —un "
        "régimen que persiguió a filósofos mientras producía a los mayores filósofos medievales de Occidente— "
        "define la complejidad de este siglo."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, este período ofrece una narrativa rica en conflictos: la expulsión de Maimónides, "
        "la construcción de la Giralda, la actividad traductora en Toledo, el pensamiento de Averroes. "
        "Son historias que conectan Al-Ándalus con la historia intelectual de toda la civilización occidental."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigación AlAndalus_SigloXII.docx")
    doc.save(out_path)
    print(f"✅ Guardado: {out_path}")


# ═════════════════════════════════════════════════════════════════════════════
# SIGLO XIII
# ═════════════════════════════════════════════════════════════════════════════

def create_siglo_XIII():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    build_cover(doc,
        "AL-ÁNDALUS EN EL SIGLO XIII",
        "El Gran Derrumbe Almohade y el Nacimiento del Reino de Granada",
        "c. 1200 – 1299")

    # ── INTRODUCCIÓN ──────────────────────────────────────────────────────────
    set_heading(doc, "Introducción", 1)
    add_paragraph(doc, (
        "El siglo XIII es, probablemente, el más dramático de toda la historia de Al-Ándalus. "
        "En pocas décadas, la catástrofe de las Navas de Tolosa (1212) derrumbó el poder almohade, "
        "desencadenando un proceso de conquistas cristianas que redujo el territorio islámico peninsular "
        "a la fracción de lo que había sido. Córdoba (1236), Valencia (1238), Sevilla (1248): "
        "las grandes ciudades cayeron una tras otra. Al mismo tiempo, emergió el Reino Nazarí de Granada, "
        "último Estado islámico independiente en la Península, que sobreviviría hasta 1492."
    ))
    add_paragraph(doc, (
        "El análisis de este siglo es indispensable para el Proyecto Noor: aquí están las claves del "
        "proceso que definió los límites finales de Al-Ándalus como entidad política independiente."
    ))

    # ── 1. POLÍTICO ───────────────────────────────────────────────────────────
    set_heading(doc, "1. Ámbito Político", 1)

    set_heading(doc, "1.1 Las Navas de Tolosa (1212): el punto de inflexión", 2)
    add_paragraph(doc, (
        "La batalla de Las Navas de Tolosa supuso la mayor derrota militar del Islam peninsular. "
        "Una cruzada convocada por el papa Inocencio III reunió a los reyes de Castilla (Alfonso VIII), "
        "Aragón (Pedro II) y Navarra (Sancho VII). El ejército almohade del califa Muhammad al-Nasir "
        "fue destruido casi en su totalidad. Tras esta derrota, el Imperio Almohade entró en una crisis "
        "de legitimidad terminal: ya no podía garantizar la protección de Al-Ándalus."
    ))

    set_heading(doc, "1.2 Las terceras Taifas y la atomización (1228-1238)", 2)
    add_paragraph(doc, (
        "El colapso almohade generó un tercer ciclo de taifas, esta vez en condiciones mucho más adversas: "
        "con los reinos cristianos en posición de fuerza y sin ninguna potencia norteafricana capaz de "
        "intervenir a tiempo. En este vacío, surgieron efímeras entidades en Murcia, Niebla, Jerez y Badajoz. "
        "Algunas buscaron la alianza con Castilla o Portugal para sobrevivir, aceptando vasallajes que "
        "las convertían en protectorados de facto."
    ))

    set_heading(doc, "1.3 Las grandes conquistas cristianas", 2)
    add_paragraph(doc, (
        "La magnitud de las pérdidas territoriales en este siglo es sin precedentes:"
    ))
    items_politico = [
        "1229-1235: Jaime I de Aragón conquista Mallorca y las Baleares.",
        "1236: Fernando III de Castilla toma Córdoba — la antigua capital califal.",
        "1238: Jaime I conquista Valencia.",
        "1246: Jaén capitula ante Fernando III.",
        "1248: Fernando III conquista Sevilla después de un asedio de 16 meses.",
        "1265-1266: Murcia, sublevada, es reconquistada con ayuda del Infante Alfonso (futuro Alfonso X).",
        "1262: Alfonso X toma Jerez, Cádiz y Niebla.",
    ]
    for item in items_politico:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    set_heading(doc, "1.4 El nacimiento del Reino Nazarí de Granada (1238)", 2)
    add_paragraph(doc, (
        "En este contexto de derrumbe, Muhammad ibn Yusuf ibn Nasr —conocido como Ibn al-Ahmar— "
        "logró consolidar un principado en torno a Granada que reconoció la soberanía nominal de Castilla "
        "a cambio de su supervivencia como Estado. En 1246 firmó con Fernando III el Pacto de Jaén, "
        "comprometiéndose al pago de parias, servicio militar y cesión de ciudades. A cambio, Granada "
        "sobreviviría como reino vasallo durante dos siglos y medio más. La Alhambra comenzó a edificarse "
        "bajo su reinado."
    ))

    # ── 2. ECONÓMICO ──────────────────────────────────────────────────────────
    set_heading(doc, "2. Ámbito Económico", 1)

    set_heading(doc, "2.1 La economía nazarí emergente", 2)
    add_paragraph(doc, (
        "El Reino de Granada, aunque diminuto —unos 30.000 km² entre las actuales provincias de Granada, "
        "Málaga y Almería—, poseía una economía excepcionalmente productiva. La seda granadina era la más "
        "valorada en los mercados de Castilla y en las ferias internacionales de Champagne y Brujas. "
        "El puerto de Almería mantuvo su función exportadora, ahora con géneros que abastecían a mercados "
        "del Mediterráneo occidental controlados por genoveses y catalanes."
    ))

    set_heading(doc, "2.2 Las parias y la economía de la dependencia", 2)
    add_paragraph(doc, (
        "El pago de parias a Castilla constituía una sangría constante. Muhammad I entregaba anualmente "
        "150.000 maravedíes de oro, cifra que representaba una proporción significativa de los ingresos "
        "fiscales del reino. Esto obligó a una fiscalidad interior intensa, que recaía sobre la producción "
        "sedera y el comercio, pero que también generaba resentimiento social."
    ))

    set_heading(doc, "2.3 Las rutas comerciales post-conquista", 2)
    add_paragraph(doc, (
        "La pérdida de Sevilla y Valencia desarticuló las antiguas redes comerciales andalusíes. "
        "Granada reorientó su comercio hacia el norte de África, especialmente hacia los marinidas "
        "de Marruecos, que se convirtieron en sus principales aliados diplomáticos y militares. "
        "Esta conexión transestrechal fue básica para la supervivencia del reino."
    ))

    # ── 3. SOCIAL ─────────────────────────────────────────────────────────────
    set_heading(doc, "3. Ámbito Social", 1)

    set_heading(doc, "3.1 Las migraciones masivas: el éxodo andalusí", 2)
    add_paragraph(doc, (
        "Las conquistas del siglo XIII desencadenaron la mayor migración forzada de la historia "
        "medieval hispana. Centenares de miles de musulmanes de Córdoba, Valencia, Sevilla y otras "
        "ciudades conquistadas huyeron o fueron expulsados. Muchos se refugiaron en Granada, que vivió "
        "una explosión demográfica. Otros cruzaron el Estrecho hacia el Magreb, llevando consigo técnicas "
        "agrícolas, artesanales y culturales que transformaron el norte de África."
    ))

    set_heading(doc, "3.2 Los mudéjares: musulmanes bajo dominio cristiano", 2)
    add_paragraph(doc, (
        "No todos los musulmanes huyeron. Amplias comunidades —los mudéjares— permanecieron en los "
        "territorios conquistados, sometidas a un estatuto jurídico de tolerancia condicionada. "
        "Eran imprescindibles como artesanos, agricultores y especialistas en sistemas hidráulicos. "
        "En Valencia, los mudéjares constituyeron un tercio de la población durante generaciones. "
        "Su aportación al arte mudéjar —síntesis de técnicas islámicas con estructuras de encargo cristiano— "
        "fue una de las mayores contribuciones culturales del período."
    ))

    set_heading(doc, "3.3 La sociedad granadina: un crisol de refugiados", 2)
    add_paragraph(doc, (
        "Granada se convirtió en receptor de oleadas sucesivas de refugiados procedentes de "
        "Andalucía, Murcia, Valencia y las Baleares. Esta heterogeneidad de orígenes enriqueció "
        "culturalmente al reino pero también generó tensiones entre grupos regionales. "
        "La élite dirigente nazarí tuvo que construir una identidad colectiva que "
        "transcendiera las lealtades tribales y regionales de los recién llegados."
    ))

    # ── 4. CULTURAL ───────────────────────────────────────────────────────────
    set_heading(doc, "4. Ámbito Cultural y Científico", 1)

    set_heading(doc, "4.1 El inicio de la Alhambra", 2)
    add_paragraph(doc, (
        "Muhammad I ordenó las primeras obras en la colina de la Sabika que se convertirían en la Alhambra. "
        "Aunque el gran desarrollo del palacio pertenece al siglo XIV, las estructuras defensivas y residenciales "
        "iniciales se levantaron en el siglo XIII. La elección del emplazamiento era simbólica: desde allí "
        "se dominaba visualmente toda la vega granadina y se transmitía un mensaje de permanencia y legitimidad."
    ))

    set_heading(doc, "4.2 Alfonso X el Sabio y la herencia andalusí", 2)
    add_paragraph(doc, (
        "Paradójicamente, uno de los grandes herederos culturales de Al-Ándalus fue el rey castellano "
        "Alfonso X el Sabio (1252-1284). Su corte en Sevilla —recién conquistada— acogió a intelectuales "
        "musulmanes y judíos que tradujeron textos árabes y contribuyeron a obras como las Tablas Alfonsinas "
        "astronómicas, el Libro del ajedrez y las Cantigas de Santa María. Este flujo cultural "
        "demuestra la permeabilidad de las fronteras intelectuales incluso en plena Reconquista."
    ))

    set_heading(doc, "4.3 Ibn al-Abbar y la literatura del exilio", 2)
    add_paragraph(doc, (
        "El valenciano Ibn al-Abbar (1199-1260) encarna la figura del intelectual andalusí exiliado. "
        "Secretario de la cancillería almohade, tras la caída de Valencia marchó a Túnez donde continuó "
        "su obra historiográfica y poética. Su Hulla assiyara es una antología literaria de enorme valor. "
        "Esta «literatura del exilio» andalusí —nostálgica, elegíaca— es uno de los corpus más conmovedores "
        "de la poesía árabe medieval."
    ))

    set_heading(doc, "4.4 Ciencias: astronomía y medicina", 2)
    add_paragraph(doc, (
        "Las tradiciones científicas andalusíes continuaron en el siglo XIII, aunque desplazadas "
        "geográficamente. Ibn al-Baitar (Málaga, c. 1190 – Damasco, 1248), el mayor botánico medieval, "
        "recorrió toda la cuenca mediterránea catalogando más de 1.400 plantas medicinales en su Tratado "
        "de los simples. Ibn Sabin de Murcia (1217-1269) mantuvo correspondencia filosófica con el "
        "Emperador Federico II de Hohenstaufen —las llamadas Cuestiones Sicilianas— demostrando "
        "la vigencia del diálogo intelectual transpirenaico."
    ))

    # ── 5. BÚSQUEDA BIBLIOGRÁFICA ─────────────────────────────────────────────
    set_heading(doc, "5. Plan de Búsqueda Bibliográfica", 1)
    add_paragraph(doc, (
        "El siguiente plan sistematizado permite recuperar la producción académica en WOS y Scopus "
        "para el análisis multidimensional del siglo XIII en Al-Ándalus."
    ))
    create_search_table(doc, [
        ("1. Político", "WOS",
         '("Al-Andalus" OR "Islamic Iberia") AND ("Navas de Tolosa" OR "Nasrid*" OR "Reconquista" OR "Fernando III" OR "Almohad*") AND ("13th century" OR "thirteenth century") AND PY=(1980-2025)'),
        ("2. Económico", "WOS/Scopus",
         '("Granada" OR "Nasrid") AND ("silk" OR "parias" OR "trade" OR "Marinid" OR "Mediterranean") AND ("medieval" OR "13th century")'),
        ("3. Social", "WOS",
         '("Al-Andalus" OR "Islamic Iberia") AND ("Mudéjar*" OR "mudajar*" OR "refugee*" OR "migration" OR "Muslim minority") AND ("13th century" OR "thirteenth century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Nasrid" OR "Granada") AND ("Alhambra" OR "Alfonso X" OR "Ibn al-Baitar" OR "Ibn al-Abbar" OR "translation" OR "poetry")'),
        ("5. Transversal", "Ambas",
         'TS=("Al-Andalus" OR "Nasrid*" OR "Islamic Iberia" OR "Granada kingdom") AND PY=(1980-2025) NOT TS=("Ottoman" OR "Crusades" OR "Egypt")'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodológicas", 2)
    items = [
        "Exportar resultados en formato BibTeX o RIS hacia Zotero o Mendeley para deduplicar entre bases de datos.",
        "Revisar directamente: Al-Qantara (CSIC), Journal of Medieval Iberian Studies, Arabica, Medieval Encounters.",
        "Para estudios sobre mudéjares: base de datos Mudéjar (CSIC) y Actas del Simposio Internacional de Mudejarismo.",
        "Consultar el corpus The Book of the Fundamento de la Astronomía para fuentes alfonsíes.",
        "Filtro de idiomas: EN, ES, FR, AR. Periodo preferente: 1990-2025 (sesgo hacia estudios más recientes).",
        "Para el exilio andalusí en el Magreb: buscar en bibliotecas de manuscritos de Túnez, Rabat y Fez.",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    # ── CONCLUSIÓN ────────────────────────────────────────────────────────────
    set_heading(doc, "Conclusión", 1)
    add_paragraph(doc, (
        "El siglo XIII enseña que los cataclismos históricos no destruyen la civilización: la reconfiguran. "
        "Al-Ándalus no murió en Las Navas de Tolosa ni con la caída de Sevilla. Se concentró, se depuró "
        "y se reinventó en el último reducto granadino. Al mismo tiempo, sus conocimientos, técnicas y "
        "sensibilidad estética se dispersaron por el Mediterráneo, enriqueciendo tanto al mundo cristiano "
        "como al islámico."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, este siglo ofrece las historias más dramáticas del ciclo: el éxodo, "
        "la supervivencia de Granada, la Alhambra naciente, Alfonso X heredando la sabiduría de sus "
        "enemigos. Son historias de pérdida y de adaptación que conectan con profundidad con el público "
        "contemporáneo."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigación AlAndalus_SigloXIII.docx")
    doc.save(out_path)
    print(f"✅ Guardado: {out_path}")


# ═════════════════════════════════════════════════════════════════════════════
# SIGLO XIV
# ═════════════════════════════════════════════════════════════════════════════

def create_siglo_XIV():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    build_cover(doc,
        "AL-ÁNDALUS EN EL SIGLO XIV",
        "El Esplendor Nazarí: La Alhambra y el Último Florecimiento",
        "c. 1300 – 1399")

    # ── INTRODUCCIÓN ──────────────────────────────────────────────────────────
    set_heading(doc, "Introducción", 1)
    add_paragraph(doc, (
        "El siglo XIV representa la cumbre y la crisis simultánea del último Estado islámico independiente "
        "de la Península Ibérica: el Reino Nazarí de Granada. Es el siglo de la Alhambra, del gran esplendor "
        "de la corte granadina y de algunos de los más refinados palacios medievales del mundo. "
        "Pero es también el siglo de las guerras civiles dinásticas, de la Peste Negra, de las batallas "
        "del Salado y del Río Palmones, y de una vulnerabilidad estructural que anunciaba el desenlace final."
    ))
    add_paragraph(doc, (
        "Este informe analiza el siglo XIV en Al-Ándalus bajo las cuatro dimensiones del Proyecto Noor "
        "—política, económica, social y cultural—, con el consiguiente plan de búsqueda bibliográfica "
        "en Web of Science y Scopus."
    ))

    # ── 1. POLÍTICO ───────────────────────────────────────────────────────────
    set_heading(doc, "1. Ámbito Político", 1)

    set_heading(doc, "1.1 La Alhambra como proyecto de legitimación", 2)
    add_paragraph(doc, (
        "Los sultanes Muhammad II (1273-1302), Muhammad III (1302-1309) e Ismail I (1314-1325) impulsaron "
        "las primeras grandes fases constructivas de la Alhambra como proyecto político-simbólico. "
        "El palacio no era solo residencia: era una declaración de soberanía, un escenario de poder "
        "diseñado para impresionar a embajadores, vasallos y súbditos. La epigrafía coránica que recorre "
        "sus muros proclama la grandeza de Dios y, por extensión, la del sultanato que construye en Su nombre."
    ))

    set_heading(doc, "1.2 Las guerras civiles nazaríes", 2)
    add_paragraph(doc, (
        "El siglo XIV fue políticamente inestable para Granada. La sucesión sultánica fue frecuentemente "
        "violenta: varios sultanes fueron depuestos o asesinados en conspiraciones palaciegas. "
        "Las facciones tribales bereberes, especialmente los Abencerrajes y los Zegríes, protagonizaron "
        "luchas por el control de la corte. Esta inestabilidad interna debilitó la capacidad defensiva "
        "del reino en los momentos críticos."
    ))

    set_heading(doc, "1.3 La amenaza marinida y la batalla del Salado (1340)", 2)
    add_paragraph(doc, (
        "La alianza entre Granada y los Marinidas de Marruecos fue el recurso estratégico básico del "
        "reino para compensar su inferioridad frente a Castilla. En 1340, un gran ejército combinado "
        "granadino-marinida fue derrotado por Alfonso XI de Castilla y Alfonso IV de Portugal en la "
        "batalla del Salado (Tarifa). Esta derrota eliminó definitivamente la posibilidad de una "
        "intervención norteafricana a gran escala y dejó a Granada sin su principal apoyo exterior."
    ))

    set_heading(doc, "1.4 Muhammad V y el esplendor diplomático (1354-1391)", 2)
    add_paragraph(doc, (
        "El reinado de Muhammad V fue el más largo y brillante del siglo. Restituido en el trono "
        "tras una deposición temporal, gobernó hasta 1391 con habilidad diplomática excepcional. "
        "Mantuvo relaciones con Castilla, Francia, el Papado y el sultán de Egipto. Bajo su mecenazgo "
        "se construyeron los palacios más refinados de la Alhambra: el Palacio de Comares y el Palacio de "
        "los Leones, inaugurado en 1362. Su corte fue el escenario donde Ibn Jaldún, el mayor historiador "
        "medieval, visitó Granada y reflexionó sobre el ocaso de las civilizaciones."
    ))

    # ── 2. ECONÓMICO ──────────────────────────────────────────────────────────
    set_heading(doc, "2. Ámbito Económico", 1)

    set_heading(doc, "2.1 La economía de la seda", 2)
    add_paragraph(doc, (
        "La seda fue el eje de la economía granadina del siglo XIV. La Alpujarra —la franja montañosa "
        "entre Sierra Nevada y el mar— albergaba miles de moreras y miles de familias especializadas "
        "en la cría del gusano de seda. El producto era exportado en bruto o trabajado en los talleres "
        "de Granada, cuyas telas eran vendidas en las ferias de Castilla, en Italia y en el Levante. "
        "La alcaicería de Granada, barrio comercial especializado en la seda, era el corazón económico "
        "de la ciudad."
    ))

    set_heading(doc, "2.2 La Peste Negra y sus consecuencias económicas (1348-1350)", 2)
    add_paragraph(doc, (
        "La Peste Negra azotó Al-Ándalus con una virulencia comparable a la que devastó el resto de "
        "Europa. Las estimaciones más recientes sugieren una mortalidad del 30-40% en las ciudades "
        "granadinas. Las consecuencias económicas fueron severas: caída de la producción sedera, "
        "escasez de mano de obra, encarecimiento de los salarios agrícolas y dislocación de las "
        "redes comerciales. Sin embargo, Granada se recuperó más rápidamente que otros territorios "
        "gracias a la resiliencia de su sistema agrícola intensivo."
    ))

    set_heading(doc, "2.3 Comercio y relaciones con Castilla y Génova", 2)
    add_paragraph(doc, (
        "El siglo XIV vio la paradoja de una intensa interdependencia económica entre Granada y sus "
        "adversarios políticos. Los mercaderes genoveses, que ya controlaban el comercio de exportación "
        "andalusí, establecieron factorías en Almería y Málaga. Castilla compraba seda granadina y "
        "vendía cereales al reino, cuya agricultura de montaña no producía suficiente grano. "
        "Esta interdependencia económica moderó en ocasiones la presión militar."
    ))

    # ── 3. SOCIAL ─────────────────────────────────────────────────────────────
    set_heading(doc, "3. Ámbito Social", 1)

    set_heading(doc, "3.1 La sociedad granadina madura", 2)
    add_paragraph(doc, (
        "El siglo XIV vio la consolidación de una sociedad granadina específica, distinta de las "
        "anteriores formaciones andalusíes. Era una sociedad homogéneamente musulmana —sin minorías "
        "cristiana y con judíos reducidos a un papel marginal—, pero internamente diversa por la "
        "procedencia de sus habitantes: descendientes de refugiados de diferentes regiones de la "
        "antigua Al-Ándalus, bereberes de distintas tribus, mercenarios marinidas y esclavos de "
        "procedencia subsahariana y eslava."
    ))

    set_heading(doc, "3.2 Los ulemas y el poder religioso", 2)
    add_paragraph(doc, (
        "El cuerpo de ulemas —juristas, teólogos y líderes religiosos— jugó en Granada un papel "
        "político crucial. Su apoyo era indispensable para la legitimación de los sultanes; su oposición "
        "podía tumbar gobiernos. El siglo XIV vio un creciente conservadurismo jurídico que limitaba "
        "los márgenes de experimentación intelectual. Sin embargo, la tradición sufí también tenía raíces "
        "profundas en la sociedad granadina, con figuras de gran influencia popular."
    ))

    set_heading(doc, "3.3 Los cautivos y el intercambio de prisioneros", 2)
    add_paragraph(doc, (
        "La frontera con Castilla generó un régimen específico de intercambio de cautivos, rescates "
        "y conversiones. Existían instituciones especializadas en el rescate de prisioneros: alfaqueques "
        "(mediadores profesionales de ambos lados de la frontera). Esta actividad era económicamente "
        "significativa y socialmente compleja: muchos cautivos pasaban años en el reino enemigo antes "
        "de ser rescatados o de adaptarse a su nueva condición."
    ))

    # ── 4. CULTURAL ───────────────────────────────────────────────────────────
    set_heading(doc, "4. Ámbito Cultural y Científico", 1)

    set_heading(doc, "4.1 La Alhambra: cumbre del arte islámico medieval", 2)
    add_paragraph(doc, (
        "El Palacio de los Leones, construido bajo Muhammad V entre 1354 y 1391, es la obra cumbre "
        "del arte nazarí y uno de los edificios más influyentes de la historia de la arquitectura. "
        "Su patio, con la famosa fuente de los doce leones de mármol, sus salas de mocárabes y su "
        "decoración de yeso tallado representan la síntesis más refinada de los lenguajes ornamentales "
        "del Islam peninsular. La epigrafía de los poemas de Ibn Zamrak incrustada en sus muros convierte "
        "el edificio en un poema arquitectónico."
    ))

    set_heading(doc, "4.2 Ibn Jaldún en Granada y la teoría de la civilización", 2)
    add_paragraph(doc, (
        "Ibn Jaldún (Túnez, 1332 – El Cairo, 1406) visitó Granada en 1363, donde fue recibido "
        "por Muhammad V con todos los honores. Su obra capital, la Muqaddima (Prolegómenos), escrita "
        "en 1377, es el primer intento sistemático de una teoría científica de la historia y la sociedad. "
        "Observando el declive de las civilizaciones como la nazarí, desarrolló su concepto de asabiyya "
        "(cohesión social) como motor del ciclo histórico. Es el padre fundador de la sociología y la "
        "historiografía crítica."
    ))

    set_heading(doc, "4.3 Poesía: Ibn Zamrak y la corte granadina", 2)
    add_paragraph(doc, (
        "Ibn Zamrak (Granada, 1333-1393) fue el poeta oficial de la corte nazarí y el autor de los poemas "
        "inscritos en los muros de la Alhambra. Sus versos, integrados en la arquitectura, describen "
        "los propios palacios desde dentro, convirtiendo la experiencia de habitarlos en una experiencia "
        "literaria simultánea. Su vida también ejemplifica el riesgo del intelectual cortesano: fue "
        "ejecutado por orden del sultán Muhammad VII."
    ))

    set_heading(doc, "4.4 Medicina, botánica y ciencias aplicadas", 2)
    add_paragraph(doc, (
        "Ibn al-Khatib (Loja, 1313 – Fez, 1374), polígrafo, médico y estadista, es una de las "
        "figuras más fascinantes del siglo. Como secretario y visir de Muhammad V, escribió obras "
        "de historia, geografía, medicina y poesía. Su texto sobre la Peste Negra fue uno de los "
        "primeros en proponer una teoría de contagio persona a persona, anticipando conceptos "
        "epidemiológicos modernos. Murió asesinado víctima de intrigas políticas en Fez."
    ))

    # ── 5. BÚSQUEDA BIBLIOGRÁFICA ─────────────────────────────────────────────
    set_heading(doc, "5. Plan de Búsqueda Bibliográfica", 1)
    add_paragraph(doc, (
        "El siguiente plan permite recuperar la producción académica especializada en WOS y Scopus "
        "para el análisis multidimensional del siglo XIV nazarí."
    ))
    create_search_table(doc, [
        ("1. Político", "WOS",
         '("Al-Andalus" OR "Nasrid" OR "Granada") AND ("Alhambra" OR "Muhammad V" OR "Battle of Salado" OR "Marinid" OR "dynastic") AND ("14th century" OR "fourteenth century") AND PY=(1980-2025)'),
        ("2. Económico", "WOS/Scopus",
         '("Nasrid" OR "Granada") AND ("silk" OR "trade" OR "plague" OR "Black Death" OR "Genoese" OR "economy") AND ("medieval" OR "14th century")'),
        ("3. Social", "WOS",
         '("Al-Andalus" OR "Nasrid Granada") AND ("ulema*" OR "frontier" OR "captive*" OR "alfaqueque*" OR "social structure") AND ("14th century" OR "fourteenth century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Nasrid" OR "Granada") AND ("Ibn Khaldun" OR "Ibn Zamrak" OR "Alhambra" OR "Ibn al-Khatib" OR "muqarnas" OR "architecture")'),
        ("5. Transversal", "Ambas",
         'TS=("Nasrid*" OR "Al-Andalus" OR "Granada kingdom" OR "Islamic Iberia") AND PY=(1980-2025) NOT TS=("Ottoman" OR "early modern" OR "16th century")'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodológicas", 2)
    items = [
        "Exportar resultados en formato BibTeX o RIS hacia Zotero o Mendeley para deduplicar entre bases de datos.",
        "Revisar directamente: Al-Qantara (CSIC), Journal of Medieval Iberian Studies, Arabica, Muqarnas (MIT Press).",
        "Para la Alhambra: consultar el catálogo del Patronato de la Alhambra y Generalife (Granada) y Cuadernos de la Alhambra.",
        "Para Ibn Jaldún: buscar en Annales: Histoire, Sciences Sociales y Arabic Sciences and Philosophy.",
        "Para la Peste Negra: cruzar con bases de datos de historia de la medicina y epidemiología histórica.",
        "Filtro temporal recomendado: 1985-2025. Idiomas: EN, ES, FR, AR.",
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)

    # ── CONCLUSIÓN ────────────────────────────────────────────────────────────
    set_heading(doc, "Conclusión", 1)
    add_paragraph(doc, (
        "El siglo XIV nazarí es la historia de una civilización que alcanza su cénit artístico cuando "
        "su horizonte político se estrecha irremisiblemente. La Alhambra no es un monumento triunfal: "
        "es la obra de un Estado que sabe que puede desaparecer y que elige, en ese contexto, "
        "crear algo de una belleza inmortal. Ibn Jaldún, observando Granada, entendió que las "
        "civilizaciones tienen ciclos inevitables; Ibn al-Khatib entendió el contagio antes de que "
        "existiera la microbiología; Ibn Zamrak convirtió la arquitectura en literatura."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XIV ofrece la historia más visualmente poderosa de toda "
        "la serie: la Alhambra, la corte de Muhammad V, el pensamiento de Ibn Jaldún, la muerte "
        "heroica de Ibn al-Khatib. Son figuras e imágenes que resuenan en el presente con una "
        "inmediatez extraordinaria."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigación AlAndalus_SigloXIV.docx")
    doc.save(out_path)
    print(f"✅ Guardado: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    print("Generando Planes de Investigacion Al-Andalus - Proyecto Noor")
    print("=" * 60)
    create_siglo_XII()
    create_siglo_XIII()
    create_siglo_XIV()
    print("=" * 60)
    print("COMPLETADO. Los tres documentos han sido guardados en:")
    print(OUTPUT_DIR)
