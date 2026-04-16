# -*- coding: utf-8 -*-
"""
Generador de Plan de Investigacion Al-Andalus: Siglo X (900-999)
El Califato de Cordoba: Apogeo de la Civilizacion Andalusi
Proyecto Noor - Investigacion Historica Multidimensional
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

OUTPUT_DIR = r"C:\Users\leoga\Desktop\Noor\Carpeta Plan de invest"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
            run.font.size = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_table_row(table, col1, col2, col3, bold=False, bg_color=None):
    row = table.add_row()
    cells = row.cells
    for i, text in enumerate([col1, col2, col3]):
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if bold:
                    run.bold = True
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            cells[i]._tc.get_or_add_tcPr().append(shading)
    return row

def create_search_table(doc, rows_data):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Bloque', 'Base', 'Cadena de busqueda']
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A376C')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for bloque, base, cadena in rows_data:
        add_table_row(table, bloque, base, cadena)
    doc.add_paragraph()

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

def build_cover(doc, title, subtitle, periodo):
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
    p2 = doc.add_paragraph(periodo)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(12)
    p3 = doc.add_paragraph("Informe historico-bibliografico multidimensional")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.font.size = Pt(11)
        run.italic = True
    p4 = doc.add_paragraph("Proyecto Noor - Elaborado con asistencia de Claude (Anthropic) - 2026")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


# ===========================================================================
# SIGLO X (900-999): EL CALIFATO DE CORDOBA
# ===========================================================================

def create_siglo_X():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO X",
        "El Califato de Cordoba: Cumbre de la Civilizacion Andalusi",
        "c. 900 - 999")

    # ── INTRODUCCION ──────────────────────────────────────────────────────────
    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "El siglo X representa el cenit indiscutible de Al-Andalus como proyecto civilizatorio. "
        "En el ano 929, Abd al-Rahman III proclamo el califato independiente de Cordoba, "
        "rompiendo formalmente con la autoridad nominal de los abbasies de Bagdad y situando "
        "a la Peninsula Iberica en el epicentro de la politica mediterranea. Durante las decadas "
        "siguientes, Cordoba se convirtio en la ciudad mas poblada de Europa occidental, "
        "con estimaciones que oscilan entre 100.000 y 500.000 habitantes segun la fuente consultada. "
        "La Gran Mezquita de Cordoba recibio sus ampliaciones mas fastuosas, y el complejo "
        "palatino-administrativo de Medina Azahara (Madinat al-Zahra) se construyo como "
        "simbolo del poder y esplendor califal."
    ))
    add_paragraph(doc, (
        "El califato de 'Abd al-Rahman III (912-961) y su hijo al-Hakam II (961-976) "
        "representaron la doble cara de la grandeza andalusi: el primero fue el consolidador "
        "militar y diplomatico que sometio a los reinos del norte y establecia relaciones "
        "con Bizancio, el Sacro Imperio Romano y el norte de Africa; el segundo fue el "
        "mecenas intelectual que convirtio la biblioteca de Cordoba en una de las mas grandes "
        "del mundo mediterraneo, con centenares de miles de volumenes catalogados."
    ))
    add_paragraph(doc, (
        "La ultima decada del siglo, dominada por el hayib (primer ministro) al-Mansur ibn "
        "Abi Amir —'Almanzor' en las cronicas latinas—, ofrece el contraste dramatico del "
        "poder sin titulo: un hombre que goberno Al-Andalus con mano de hierro, humillo a "
        "los califas hasjimies y lanzo mas de cincuenta campanas militares victoriosas contra "
        "los reinos cristianos, pero cuya ambicion sembrara las semillas de la fitna "
        "(guerra civil) que destruiria el propio califato en el siglo siguiente. "
        "Para el Proyecto Noor, el siglo X es el capitulo del esplendor: la demostracion "
        "empirica de que Al-Andalus alcanzo cotas de sofisticacion intelectual, artistica "
        "y administrativa que situaron a la civilizacion islamica ibérica a la vanguardia del mundo medieval."
    ))

    # ── 1. AMBITO POLITICO ────────────────────────────────────────────────────
    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 La proclamacion del Califato (929) y su significacion historica", 2)
    add_paragraph(doc, (
        "En el ano 300 de la Hegira (912 d.C.), Abd al-Rahman III, con apenas veintiun anos, "
        "heredo un emirato fragmentado por las rebeliones de las marcas —en especial la de "
        "Umar ibn Hafsun en Bobastro— y sometido a incursiones de fatimies y leoneses. "
        "En los primeros anos de su reinado, Abd al-Rahman III consolido militarmente el "
        "territorio mediante campanas sistematicas (aceifas de verano, razias de invierno) "
        "que sometieron progresivamente a los senores rebeldes del interior y re-establecieron "
        "la soberania omeya sobre las ciudades del Magreb septentrional."
    ))
    add_paragraph(doc, (
        "La proclamacion del califato en el ano 929 fue un acto de ruptura deliberada: "
        "Abd al-Rahman III adopto los titulos de Amir al-Mu'minin (Principe de los Creyentes) "
        "y al-Nasir li-Din Allah (el Victorioso por la Religion de Dios), rechazando la "
        "autoridad nominal de los abbasies de Bagdad y desafiando al mismo tiempo a los "
        "fatimies shies de Ifriqiya. El uso del titulo califal no era meramente simbolico: "
        "afirmaba la plena soberania religiosa y politica de Cordoba y la independencia "
        "de una tradicion sunni omeya que se presentaba como genuina heredera del Islam primitivo."
    ))

    set_heading(doc, "1.2 Las relaciones diplomaticas del califato cordobes", 2)
    add_paragraph(doc, (
        "Una de las caracteristicas mas notables del califato de Abd al-Rahman III fue la "
        "intensidad de su actividad diplomatica. Embajadas procedentes de Bizancio (Constantino VII "
        "Porfirogeneta), del Sacro Imperio Romano (Otto I), del papa Juan XII, de los reinos "
        "de Leon, Navarra y los condados catalanes, de los reinos germanos y de los principes "
        "del norte de Africa llegaron a Cordoba o fueron enviadas desde ella. "
        "Esta actividad diplomatica, detalladamente registrada en las cronicas arabes —especialmente "
        "en la obra de Ibn Hayyan—, es fuente fundamental para comprender la posicion "
        "de Al-Andalus en el sistema de relaciones internacionales del siglo X."
    ))
    add_paragraph(doc, (
        "La relacion con los reinos del norte fue una mezcla de presion militar y negociacion "
        "tacticamente calibrada. Abd al-Rahman III nunca intento la conquista definitiva "
        "de los reinos cristianos —Castilla, Leon, Navarra, los condados catalanes— sino "
        "su sometimiento nominal como vasallos que pagaban parias (tributos) o solicitaban "
        "la mediacion califal en sus conflictos internos. Este sistema de supremacia tactica "
        "sin anexion territorial fue una de las grandes innovaciones diplomaticas del califato."
    ))

    set_heading(doc, "1.3 Al-Hakam II: el califa bibliofilo y la plenitud del califato", 2)
    add_paragraph(doc, (
        "Al-Hakam II (961-976) heredo de su padre un califato solido y lo convirtio en "
        "un proyecto intelectual de primer orden. Su obsesion por los libros —habia financiado "
        "agentes en Bagdad, El Cairo, Damasco y Alejandria para adquirir manuscritos— resulto "
        "en una biblioteca que las fuentes arabes cifran en centenares de miles de volumenes "
        "(las estimaciones van de 40.000 a 400.000 segun la fuente). Bajo su reinado, Cordoba "
        "atrajo a los mejores medicos, filosofos, matematicos, astronomos y poetas del "
        "mundo islamico."
    ))
    add_paragraph(doc, (
        "Politicamente, al-Hakam II continuo la politica de su padre: campanas periodicas "
        "contra los reinos del norte (con la presencia personal del califa en varias de ellas), "
        "mantenimiento del sistema de parias y vigilancia sobre las fronteras magrebies. "
        "Su reinado fue esencialmente pacifico en el interior, lo que permitio el florecimiento "
        "cultural. La Gran Mezquita de Cordoba recibio durante su reinado la "
        "ampliacion mas sofisticada: la maqsura y el mihrab, con sus mosaicos de teselas "
        "procedentes de Constantinopla, son joyas del arte islamico universal."
    ))

    set_heading(doc, "1.4 Almanzor y el poder sin titulo (976-1002)", 2)
    add_paragraph(doc, (
        "La minoria de edad de Hisham II (976) abrio el camino al ascenso de Muhammad ibn "
        "Abi Amir, conocido como al-Mansur (el Victorioso), 'Almanzor' en las fuentes latinas. "
        "Administrador brillante y general excepcional, Almanzor concentro progresivamente "
        "el poder efectivo desplazando a la familia califal de las decisiones reales. "
        "Entre 977 y 1002, dirigio mas de cincuenta campanas militares contra los reinos "
        "del norte, alcanzando victorias que ninguno de sus predecesores habia logrado: "
        "Zamora (981), Leon (988), el monasterio de San Millan de la Cogolla (994) "
        "y, en su campaña mas famosa y simbolicamente devastadora, Santiago de Compostela (997), "
        "cuyos campanarios hizo transportar en hombros de prisioneros a Cordoba "
        "para ser convertidos en lamparas de la Gran Mezquita."
    ))
    add_paragraph(doc, (
        "Almanzor fue tambien un reformador administrativo: reorganizo el ejercito "
        "introduciendo contingentes de bereberes del norte de Africa para reducir su "
        "dependencia de las tropas eslavas (saqaliba) y de los arabe-andalusies cuya "
        "lealtad era dudosa. Esta militarizacion del califato y la concentracion de poder "
        "en sus manos sentaron, paradojicamente, las bases de la fitna que estallaria "
        "tras su muerte en 1002 y que destruiria el califato en pocos decenios."
    ))

    # ── 2. AMBITO ECONOMICO ───────────────────────────────────────────────────
    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 La economia califal: produccion, circulacion y fiscalidad", 2)
    add_paragraph(doc, (
        "El califato de Cordoba del siglo X fue una potencia economica de primer orden en "
        "el mundo mediterraneo. Las fuentes arabes y los datos arqueologicos convergen en "
        "describir una economia monetizada, con una produccion artesanal diversificada, "
        "una agricultura intensiva de regadio y unas redes de comercio de largo alcance. "
        "La ceca de Cordoba —la Casa de la Moneda— acunaba dirhems de plata de alta pureza "
        "que circulaban desde el norte de Europa (las cronicas escandinavas mencionan moneda "
        "andalusi) hasta el Africa subsahariana."
    ))
    add_paragraph(doc, (
        "La base fiscal del califato descansaba sobre el impuesto sobre la tierra (jaray), "
        "los derechos de aduana (portazgos), la jizya pagada por dhimmis (judios y cristianos), "
        "los beneficios del comercio exterior y el bot in de las campanas militares. "
        "Abd al-Rahman III duplico los ingresos del Estado mediante una racionalizacion "
        "administrativa que implicaba el control directo de las principales fuentes de riqueza "
        "por parte de la hacienda califal."
    ))

    set_heading(doc, "2.2 La agricultura andalusi: sistemas de riego e innovacion agronomica", 2)
    add_paragraph(doc, (
        "Uno de los legados mas duraderos del califato fue el desarrollo de sistemas de "
        "irrigacion sofisticados que transformaron el paisaje agrario de la Peninsula. "
        "Las norias de elevacion de agua, las acequias derivadas de rios, los sistemas de "
        "noria y shaduf permitieron el cultivo de productos que requieren riego controlado: "
        "naranja amarga, limon, alubia, berenjena, alcachofa, esparrago, azucar de cana "
        "y numerosas especias. El agronomo ibn Bassal y los tratados del Calendario de "
        "Cordoba (siglo X) documentan una ciencia agronomica sistematizada que transformo "
        "radicalmente la produccion alimentaria iberica."
    ))
    add_paragraph(doc, (
        "Esta revolucion agricola andalusi —estudiada detalladamente por Andrew Watson "
        "en su obra clasica de 1983— introdujo en Europa occidental un conjunto de "
        "cultivos, tecnicas y conocimientos que fundamentaria la expansion demografica "
        "y economica de los siglos siguientes. El debate sobre la magnitud y la "
        "originalidad de esta transferencia agrontomica sigue activo en la historiografia actual."
    ))

    set_heading(doc, "2.3 El comercio de larga distancia y las rutas del oro africano", 2)
    add_paragraph(doc, (
        "Al-Andalus en el siglo X era un nodo fundamental de las rutas comerciales que "
        "conectaban el norte de Europa con el Africa subsahariana y el mundo islamico oriental. "
        "A traves del Magreb —y en particular de la ciudad de Sijilmasa en el sur de Marruecos— "
        "llegaba a Cordoba el oro del Sudan occidental (Ghana, Mali) que financiaba "
        "las emisiones de moneda, las obras de Medina Azahara y las campanas militares. "
        "En la otra direccion, los comerciantes andalusies y judios sefardies exportaban "
        "seda, cuero, ceramica y esclavos hacia el norte de Africa y el Oriente."
    ))
    add_paragraph(doc, (
        "Las comunidades de mercaderes judios —los radanitas— operaban como correistas "
        "internacionales que conectaban Al-Andalus con Bagdad, Persia y la India. "
        "Para el Proyecto Noor, esta dimension economica transnacional del califato "
        "subraya que Al-Andalus del siglo X no era un fenomeno regional sino un actor "
        "global del comercio medieval."
    ))

    set_heading(doc, "2.4 Medina Azahara: economia del lujo y poder arquitectonico", 2)
    add_paragraph(doc, (
        "La construccion de Medina Azahara (Madinat al-Zahra), iniciada por Abd al-Rahman III "
        "hacia el ano 936 y completada parcialmente bajo al-Hakam II, fue el proyecto "
        "arquitectonico mas ambicioso del mundo islamico occidental de su epoca. "
        "Las fuentes arabes mencionan cifras impresionantes: entre 10.000 y 14.000 "
        "trabajadores diarios, marmoles importados de Cartago e Italia, piezas de madera "
        "de los bosques del Atlas, columnas de la propia Cordoba clasica y relieves "
        "de piedra caliza trabajados por centenares de artesanos especializados. "
        "Los recientes trabajos arqueologicos han revelado un conjunto de mas de 112 "
        "hectareas con salones de recepcion, jardines, patios, banos termales, mezquitas "
        "y dependencias administrativas de una sofisticacion sin igual en la Europa del siglo X."
    ))

    # ── 3. AMBITO SOCIAL ──────────────────────────────────────────────────────
    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 La sociedad plural del califato: musulmanes, dhimmis y esclavos", 2)
    add_paragraph(doc, (
        "La sociedad de Al-Andalus en el siglo X era extraordinariamente compleja "
        "en su composicion. La poblacion musulmana incluia grupos de origen arabe "
        "(junds sirios, yemeníes, qaysies), bereberes del norte de Africa (llegados "
        "en oleadas desde 711), conversos ibéricos (muladies) y sus descendientes, "
        "cuya suma constituia ya en el siglo X la mayoria de la poblacion hispanica. "
        "Los dhimmis —judios y mozarabes (cristianos bajo dominio islamico)— gozaban "
        "de un estatuto juridico protegido pero subordinado, con derechos reconocidos "
        "y limitaciones especificas en la vida publica."
    ))
    add_paragraph(doc, (
        "Los judios de Al-Andalus vivieron en el siglo X uno de sus periodos mas "
        "florecientes de la diaspora. La figura de Hasdai ibn Shaprut —medico, diplomatico "
        "y mecenas de la corte califal— simboliza esta edad dorada: poliglota que hablaba "
        "arabe, hebreo, latin y las lenguas romanicas, ibn Shaprut fue el interlocutor "
        "diplomatico de Abd al-Rahman III con el Imperio Bizancio y con los jázaros, "
        "al tiempo que financiaba el estudio del Talmud y la poesia hebrea en Cordoba. "
        "La juderia de Cordoba fue en este siglo un centro intelectual de la diaspora judia "
        "comparable a Bagdad o Los Fayyum."
    ))

    set_heading(doc, "3.2 Los mozarabes y la convivencia cultural", 2)
    add_paragraph(doc, (
        "Los mozarabes —cristianos que vivian bajo dominio islamico y habian adoptado "
        "la lengua arabe como vehiculo de cultura— son uno de los fenomenos mas fascinantes "
        "de la historia ibérica. En el siglo X, las comunidades mozarabes de Cordoba, "
        "Sevilla y Toledo conservaban sus liturgias, sus iglesias y sus jerarquias "
        "eclesiasticas. Escrib ian cronicas en arabe (la denominada 'Cronica del moro Rasis' "
        "y otras), y algunos producian obras de rica hibridacion cultural, como el "
        "Beato de Liebana comentado en escritura mozarabe con ilustraciones de clara "
        "influencia islamica."
    ))
    add_paragraph(doc, (
        "El concepto historiografico de 'convivencia' —popularizado por Americo Castro "
        "y debatido intensivamente desde entonces— tiene en el siglo X uno de sus "
        "referentes empiricos mas fuertes. Sin idealizarla, la sociedad cordobesa del "
        "califato ofrecia grados de interaccion cultural, intelectual y econom ica "
        "entre grupos religiosos distintos que no tuvieron paralelo en la Europa "
        "cristiana coetanea. El debate academico sobre la naturaleza, los limites "
        "y la comparabilidad de esta convivencia es uno de los mas fecundos de los "
        "estudios medievales del siglo XXI."
    ))

    set_heading(doc, "3.3 Los saqaliba: esclavos de elite en el sistema califal", 2)
    add_paragraph(doc, (
        "Uno de los fenomenos sociales mas peculiares del califato cordobes fue el papel "
        "de los saqaliba (literalmente 'eslavos', aunque el termino designaba en realidad "
        "a cualquier esclavo de origen europeo, predominantemente eslavos y francos). "
        "Comprados en los mercados de esclavos del norte de Europa y llevados a Cordoba, "
        "los saqaliba ocupaban posiciones de confianza en el palacio: guardas personales "
        "del califa, administradores, eunucos de la corte, generales del ejercito. "
        "Su posicion entre los grandes y el pueblo comun, sin linajes tribales que "
        "proteger ni redes de lealtad propias, los convirtia en instrumentos de poder "
        "especialmente fiables para un califa que desconfiaba de las facciones arabes "
        "y bereberes."
    ))

    set_heading(doc, "3.4 La vida urbana de Cordoba: la ciudad mas grande de Europa occidental", 2)
    add_paragraph(doc, (
        "Cordoba en el siglo X era, segun las fuentes arabes y los hallazgos arqueologicos, "
        "la ciudad mas grande y compleja de Europa occidental. Las estimaciones de poblacion "
        "varian enormemente —de 100.000 a 500.000 habitantes— pero incluso las mas conservadoras "
        "situan a Cordoba muy por delante de cualquier ciudad cristiana de la epoca "
        "(Londres, Paris o Roma tenian entonces menos de 30.000 habitantes). "
        "La medina (ciudad vieja), organizada en barrios por oficios y grupos etnicos, "
        "albergaba mercados especializados, banos publicos (hammamat), mezquitas de barrio, "
        "fondas y talleres artesanales. La infraestructura urbana incluia un sistema "
        "de abastecimiento de agua por acueducto y una red de alumbrado nocturno "
        "que impresionaba a los viajeros medievales."
    ))

    # ── 4. AMBITO CULTURAL Y CIENTIFICO ───────────────────────────────────────
    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 La biblioteca de al-Hakam II y el proyecto de acumulacion del conocimiento", 2)
    add_paragraph(doc, (
        "La biblioteca del califa al-Hakam II es uno de los grandes proyectos intelectuales "
        "de la Edad Media. Las fuentes arabes afirman que contenia entre 40.000 y 400.000 "
        "volumenes —las cifras son probablemente exageradas, pero el conjunto era extraordinario "
        "para la epoca—, catalogados en 44 volumenes de indices. El califa enviaba compradores "
        "a Bagdad, Damasco, El Cairo y Alejandria para adquirir manuscritos de reciente aparicion. "
        "Cuando al-Mutanabbi —el mayor poeta arabe del siglo X— publico su primera version de "
        "un poema en Bagdad, un ejemplar llego a Cordoba antes de que el propio autor lo "
        "hubiera distribuido en el resto del mundo arabe."
    ))
    add_paragraph(doc, (
        "Esta biblioteca fue el centro gravitacional de una vida intelectual extraordinariamente "
        "rica. Filosofos como Ibn Masarra —pionero del misticismo neoplatonico que influyo "
        "en ibn 'Arabi y Averroes—, medicos como ibn Juljul —que escribio la primera "
        "historia de la medicina en Al-Andalus—, astronomos, matematicos y teoricos de la "
        "musica frecuentaban la corte y producian obras que circulaban por el mundo islamico "
        "y, en traduccion latina, por la Europa cristiana."
    ))

    set_heading(doc, "4.2 La Gran Mezquita de Cordoba: el mihrab de al-Hakam II", 2)
    add_paragraph(doc, (
        "La ampliacion de la Gran Mezquita de Cordoba bajo al-Hakam II (961-976) es "
        "la cumbre del arte arquitectonico andalusi. El nuevo mihrab —la hornacina que "
        "indica la direccion de La Meca— fue revestido con mosaicos de teselas doradas "
        "ejecutados por artesanos enviados expresamente desde Constantinopla por el "
        "emperador Niceforo II Focas en respuesta a una embajada cordobesa. El resultado "
        "es una obra de sincretismo artistico sin precedentes: tecnica bizantina, "
        "vocabulario decorativo islamico, arcos lobulados de innovacion propiamente "
        "andalusi. Junto al mihrab, la maqsura —el recinto reservado al califa— "
        "con sus arcos polilobulados entrelazados representa la maduracion plena "
        "del arco andalusi como solucion estructural y estetica."
    ))

    set_heading(doc, "4.3 La poesia y la musica en la corte califal", 2)
    add_paragraph(doc, (
        "La poesia arabe alcanzo en la corte califal cordobesa un nivel de sofisticacion "
        "extraordinario. Poetas como ibn Abd Rabbih —autor de la enciclopedia poetica "
        "'al-Iqd al-Farid' (El Collar Unico)— y la poetisa de origen esclavo Wallada "
        "bint al-Mustakfi (aunque esta florecio en el siglo XI) representan "
        "la diversidad de voces de la cultura poetica andalusi. La muwaxxaha —estrofa "
        "poetica inventada en Al-Andalus, con mezcla de arabe clasico, arabe vulgar "
        "y romance ibérico en sus finales (jarchas)— aparece ya en el siglo X como "
        "la forma lirica mas original aportada por Al-Andalus al mundo islamico."
    ))
    add_paragraph(doc, (
        "La musica de la corte —heredada de la tradicion de Ziryab, el musico bagdadi "
        "llegado a Cordoba en el siglo IX— se habia desarrollado en una tradicion propia "
        "andalusi con los modos (tabat) del sistema modal arabigo-andaluz que sobreviviria "
        "hasta hoy en las nubes (maqam) del Magreb. Al-Hakam II, segun las cronicas, "
        "era el mecenas mas generoso de musicos, y la presencia de musica de alta calidad "
        "en los banquetes de la corte era considerada signo de refinamiento civilizatorio."
    ))

    set_heading(doc, "4.4 Medicina y ciencia en el califato: Abulcasis y la escuela medica cordobesa", 2)
    add_paragraph(doc, (
        "Abu al-Qasim al-Zahrawi —Abulcasis en la tradicion latina— fue el medico y cirujano "
        "mas importante del siglo X andalusi y una de las figuras cardinales de la historia "
        "de la medicina universal. Su enciclopedia medica 'Kitab al-Tasrif' (El Libro de "
        "la Disposicion), completada hacia el 1000, contiene en su tercer volumen el "
        "primer tratado sistematico de cirugia ilustrado de la historia, con descripciones "
        "detalladas de instrumentos quirurgicos que el mismo disenaba. El Tasrif fue "
        "traducido al latin por Gerardo de Cremona en el siglo XII y utilizado como texto "
        "basico en las universidades europeas hasta el siglo XVII."
    ))
    add_paragraph(doc, (
        "En la astronomia, las tablas astronomicas (zij) elaboradas por matematicos andalusies "
        "del siglo X perfeccionaron los calculos de Ptolomeo y las contribuciones arabes "
        "orientales. El calendario agronomico de Cordoba —el 'Libro de la Anwa'— combinaba "
        "conocimiento astronomico con observacion practica de fenomenos meteorologicos "
        "y ciclos vegetales, siendo testimonio de una ciencia aplicada y enraizada "
        "en el territorio andalusi."
    ))

    set_heading(doc, "4.5 Filosofia y pensamiento: Ibn Masarra y el misticismo neoplatonico andalusi", 2)
    add_paragraph(doc, (
        "Muhammad ibn Masarra al-Qurtubi (883-931) es el primer gran filosofo originado "
        "en la Peninsula Iberica dentro de la tradicion islamica. Formado en la matematica "
        "y la filosofia griega —conocia parcialmente a Platon y los neoplatonicos a traves "
        "de traducciones arabes—, desarrollo un misticismo especulativo que mezclaba "
        "elementos del neoplatonismo con la teologia islamica. Sus ideas, consideradas "
        "heterodoxas por las autoridades religiosas de su tiempo, tuvieron que sobrevivir "
        "en circulos discretos, pero influyeron en la tradicion filosofica andalusi "
        "y, a traves de ibn 'Arabi (siglo XIII), en el sufismo universal."
    ))

    # ── 5. PLAN DE BUSQUEDA BIBLIOGRAFICA ─────────────────────────────────────
    set_heading(doc, "5. Plan de Busqueda Bibliografica Sistematizada", 1)
    add_paragraph(doc, (
        "El siguiente plan estructurado en cinco bloques tematicos permite recuperar "
        "la produccion academica relevante en Web of Science (WOS) y Scopus para el "
        "analisis historico multidimensional del siglo X de Al-Andalus y el Califato de "
        "Cordoba. Las cadenas de busqueda estan disenadas para maximizar la recuperacion "
        "en ingles, la lengua de indexacion principal en ambas bases de datos, con "
        "complemento en espanol y frances donde las bases lo permiten."
    ))
    create_search_table(doc, [
        ("1. Politico-Diplomatico", "WOS",
         '(\"Al-Andalus\" OR \"Caliphate of Cordoba\" OR \"Umayyad Caliphate\") AND '
         '(\"Abd al-Rahman III\" OR \"al-Hakam II\" OR \"Almanzor\" OR \"al-Mansur\" OR '
         '\"caliphate\" OR \"diplomacy\" OR \"political\" OR \"military campaign*\") AND '
         '(\"10th century\" OR \"tenth century\" OR \"929\" OR \"Cordoba\") AND PY=(1980-2025)'),

        ("2. Economico", "WOS/Scopus",
         '(\"Al-Andalus\" OR \"Islamic Iberia\" OR \"Cordoba\") AND '
         '(\"economy\" OR \"trade\" OR \"agriculture\" OR \"irrigation\" OR \"silk\" OR '
         '\"gold\" OR \"coinage\" OR \"dirham\" OR \"Madinat al-Zahra\" OR \"market*\") AND '
         '(\"10th century\" OR \"medieval Islamic\" OR \"Umayyad\")'),

        ("3. Social-Demografico", "WOS",
         '(\"Al-Andalus\" OR \"Caliphate of Cordoba\") AND '
         '(\"dhimmi\" OR \"mozarab\" OR \"muwallad\" OR \"saqaliba\" OR \"convivencia\" OR '
         '\"social structure\" OR \"urban life\" OR \"Jewish community\" OR \"Hasdai ibn Shaprut\") AND '
         '(\"10th century\" OR \"medieval Iberia\" OR \"Umayyad Cordoba\") AND PY=(1980-2025)'),

        ("4. Cultural-Cientifico", "WOS/Scopus",
         '(\"Al-Andalus\" OR \"Cordoba\" OR \"Andalusian\") AND '
         '(\"Abulcasis\" OR \"al-Zahrawi\" OR \"Ibn Masarra\" OR \"al-Hakam II\" OR '
         '\"Great Mosque\" OR \"architecture\" OR \"medicine\" OR \"astronomy\" OR '
         '\"poetry\" OR \"muwaxxaha\" OR \"maqama\" OR \"library\" OR \"manuscript\") AND '
         '(\"10th century\" OR \"Umayyad\" OR \"medieval Islamic science\")'),

        ("5. Transversal / Historiografico", "Ambas",
         'TS=(\"Al-Andalus\" OR \"Caliphate of Cordoba\" OR \"Umayyad Spain\" OR '
         '\"Islamic Iberia\" OR \"Madinat al-Zahra\" OR \"Mozarab*\" OR '
         '\"medieval Cordoba\") AND PY=(1980-2025) AND '
         'NOT TS=(\"modern\" OR \"contemporary\" OR \"postcolonial\")'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas avanzadas", 2)
    for item in [
        "Exportar resultados en formato BibTeX o RIS e importar directamente en Zotero (con deduplicacion automatica entre WOS y Scopus via DOI).",
        "Revistas especializadas prioritarias: Al-Qantara (CSIC, acceso digital), Journal of Medieval Iberian Studies, Arabica (Brill), Al-Masaq, Medieval Encounters, Muqarnas (arquitectura islamica).",
        "Fuentes primarias arabes digitalizadas: Biblioteca Virtual Miguel de Cervantes (fondos arabes), Bibliotheca Islamica (Brill), y el proyecto 'Manuscritos Arabes de al-Andalus' de la Real Academia de la Historia.",
        "Obras de referencia ineludibles: Mª Jesus Viguera Molins (coord.), 'El Califato Omeya de al-Andalus' (Historia de Espana / Menendez Pidal, vol. VIII-1), 1994. Pierre Guichard, 'Al-Andalus frente a la conquista cristiana'. Hugh Kennedy, 'Muslim Spain and Portugal: A Political History of al-Andalus'.",
        "Para Medina Azahara: publicaciones del Conjunto Arqueologico Madinat al-Zahra (Junta de Andalucia) y los informes de las excavaciones dirigidas por Antonio Vallejo Triano.",
        "Para Abulcasis: edicion critica de la obra de M. S. Spink y G. L. Lewis (Wellcome Institute, Londres). Para filosofia: obras de Cruz Hernandez y Dominique Urvoy.",
        "Filtro temporal recomendado: 1980-2025. Idiomas de busqueda: EN (principal), ES, FR, AR (cuando la base lo permite).",
        "Herramienta de visualizacion bibliometrica: VOSviewer aplicado a los resultados exportados para identificar clusters tematicos y autores centrales.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "5.2 Terminos de busqueda en espanol y frances (complementarios)", 2)
    add_paragraph(doc, (
        "Para las bases de datos que permiten busqueda en multiples idiomas (Scopus permite "
        "busqueda AND LANGUAGE = Spanish/French), se recomienda complementar con:"
    ))
    for item in [
        "ES: (\"Al-Andalus\" O \"Califato de Cordoba\" O \"omeya\") Y (\"siglo X\" O \"siglo decimo\" O \"Abd al-Rahman III\") Y (\"cultura\" O \"economia\" O \"politica\" O \"sociedad\")",
        "FR: (\"Al-Andalus\" OR \"Califat de Cordoue\" OR \"Omeyyade\") AND (\"Xe siecle\" OR \"dixieme siecle\") AND (\"culture\" OR \"economie\" OR \"medecine\" OR \"architecture\")",
        "AR: استخدام محركات JSTOR y Google Scholar مع مصطلحات: الأندلس - الخلافة الأموية - قرطبة - عبد الرحمن الثالث",
    ]:
        add_bullet(doc, item)

    # ── CONCLUSION ────────────────────────────────────────────────────────────
    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo X de Al-Andalus es, en terminos historicos, el periodo de mayor densidad "
        "civilizatoria de la experiencia islamica ibérica. En menos de cien anos, la Peninsula "
        "paso de ser un emirato fragmentado y amenazado a constituir el califato mas poderoso "
        "del Mediterraneo occidental, con una capital —Cordoba— que situaba a Europa "
        "en el mapa cultural y cientifico del mundo medieval de una manera que tardaria "
        "siglos en volver a producirse."
    ))
    add_paragraph(doc, (
        "Los cuatro ambitos analizados en este informe —politico, economico, social y cultural— "
        "convergen en una imagen coherente: la de una civilizacion que logro, durante un siglo "
        "excepcional, articular poder militar y diplomatico con florecimiento intelectual, "
        "pluralidad social con cohesion politica, innovacion cientifica con sofisticacion "
        "artistica. Las contradicciones existian —la esclavitud, la subordinacion de los dhimmis, "
        "la violencia de las campanas militares— pero el conjunto supera con creces la "
        "caricatura simplificadora tanto del andalucismo romantico como del revisionismo "
        "anticonvivencial."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo X ofrece el argumento narrativo central: "
        "la demostracion empirica de que Al-Andalus fue, en su momento de plenitud, "
        "una de las grandes civilizaciones de la historia humana. La figura de Abd al-Rahman III "
        "—el nino que heredo un emirato en crisis y lo convirtio en el califato mas brillante "
        "de su tiempo— y la de al-Hakam II —el califa que puso el conocimiento al servicio "
        "del bienestar de su pueblo— son los protagonistas de una historia que merece "
        "ser contada con la rigor y la emocion que el Proyecto Noor aspira a ofrecer."
    ))
    add_paragraph(doc, (
        "El legado del siglo X no termino en el ano 1000. La biblioteca de al-Hakam II, "
        "parcialmente salvada de la destruccion de Medina Azahara (saqueada en 1009-1010), "
        "alimentara durante siglos la filosofia, la medicina y la astronomia de Europa. "
        "Abulcasis seguira siendo texto universitario en Montpellier y Bolonia hasta el "
        "siglo XVII. El mihrab de la Mezquita de Cordoba seguira siendo considerado "
        "por los historiadores del arte uno de los espacios mas sagrados y bellos "
        "de la arquitectura universal. Y el concepto hispano de convivencia —con todas "
        "sus imperfecciones y sus grandezas— seguira interpelando a un mundo "
        "que todavia no ha aprendido del todo sus lecciones."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloX.docx")
    doc.save(out_path)
    print("Guardado exitosamente: " + out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 70)
    print("Generando Plan de Investigacion Al-Andalus (Siglo X) - Proyecto Noor")
    print("=" * 70)
    path = create_siglo_X()
    print("=" * 70)
    print("COMPLETADO. Documento guardado en:")
    print(path)
