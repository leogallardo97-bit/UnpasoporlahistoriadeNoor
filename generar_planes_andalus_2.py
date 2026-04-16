# -*- coding: utf-8 -*-
"""
Generador de Planes de Investigacion Al-Andalus: Siglos XV al XIX
Proyecto Noor - Investigacion Historica Multidimensional
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = r"C:\Users\leoga\Desktop\Noor\Carpeta Plan de invest"

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xA8)
            run.font.size = Pt(13)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
            run.font.size = Pt(12)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_table_row(table, col1, col2, col3, bold=False, bg_color=None):
    row = table.add_row()
    cells = row.cells
    for i, text in enumerate([col1, col2, col3]):
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                if bold:
                    run.bold = True
        if bg_color:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), bg_color)
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            cells[i]._tc.get_or_add_tcPr().append(shading)
    return row

def create_search_table(doc, rows_data):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ['Bloque', 'Base', 'Cadena de busqueda']
    for i, txt in enumerate(headers):
        hdr[i].text = txt
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1A376C')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:val'), 'clear')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for bloque, base, cadena in rows_data:
        add_table_row(table, bloque, base, cadena)
    doc.add_paragraph()

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

def build_cover(doc, title, subtitle, periodo):
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in s.runs:
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x4A, 0x80, 0xBD)
    p2 = doc.add_paragraph(periodo)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size = Pt(12)
    p3 = doc.add_paragraph("Informe historico-bibliografico multidimensional")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p3.runs:
        run.font.size = Pt(11)
        run.italic = True
    p4 = doc.add_paragraph("Proyecto Noor - Elaborado con asistencia de Claude (Anthropic) - 2025")
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p4.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()


# ===========================================================================
# SIGLO XV (1400-1499)
# ===========================================================================

def create_siglo_XV():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO XV",
        "El Ocaso Nazari: De la Pugna Interna a la Caida de Granada",
        "c. 1400 - 1499")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "El siglo XV cierra el ciclo historico de Al-Andalus como entidad politica independiente. "
        "El Reino Nazari de Granada, que habia sobrevivido dos siglos y medio gracias a una combinacion "
        "de habilidad diplomatica, solidez economica y fragmentacion entre sus enemigos castellanos y aragoneses, "
        "se derrumbo en enero de 1492 ante los ejercitos de los Reyes Catolicos. "
        "Es el siglo de las guerras civiles nazaries, de la figura tragica de Boabdil (Muhammad XII), "
        "de la rendicion de la Alhambra y del legendario 'suspiro del Moro'. "
        "Pero es tambien el siglo de una cultura granadina que siguio produciendo obras de primer orden "
        "hasta el ultimo momento."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, este siglo es un capitulo ineludible: el fin de Al-Andalus "
        "como Estado no fue el fin de su civilizacion. Las comunidades musulmanas que permanecieron "
        "en la Peninsula como mudejares y, posteriormente, como moriscos, fueron los depositarios "
        "de una memoria que tardaria siglos en apagarse."
    ))

    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 Inestabilidad dinastica y guerras civiles nazaries", 2)
    add_paragraph(doc, (
        "El siglo XV comenzó para Granada bajo el reinado de Muhammad VII (1392-1408), que mantuvo "
        "una politica de equilibrio frente a Castilla. Sin embargo, la segunda mitad del siglo "
        "fue una espiral de violencia sucesoria. Entre 1454 y 1492 se sucedieron al menos ocho sultanes, "
        "varios de ellos depuestos o asesinados. Las facciones de los Abencerrajes —partidarios de "
        "Muhammad XII (Boabdil)— y los Zegries —leales a Muhammad el Zagal, su tio— convirtieron "
        "el reino en un campo de batalla civil, facilitando enormemente la labor de los Reyes Catolicos."
    ))

    set_heading(doc, "1.2 La presion castellano-aragonesa: los Reyes Catolicos", 2)
    add_paragraph(doc, (
        "La union de las coronas de Castilla (Isabel I) y Aragon (Fernando II) en 1479 creo por "
        "primera vez una potencia capaz de dedicar todos sus recursos a la conquista final de Granada. "
        "La guerra de Granada (1482-1492) fue una campana sistematica y bien planeada: los Reyes Catolicos "
        "no buscaban el saqueo sino la incorporacion definitiva del territorio. Fueron tomando ciudades "
        "una a una —Ronda (1485), Malaga (1487), Almeria y Baza (1489)— hasta rodear completamente "
        "la capital. Boabdil, en una situacion sin salida, capituló el 2 de enero de 1492."
    ))

    set_heading(doc, "1.3 Las Capitulaciones de Granada y sus consecuencias inmediatas", 2)
    add_paragraph(doc, (
        "Las Capitulaciones de Santa Fe (noviembre de 1491) garantizaban a los musulmanes granadinos "
        "el ejercicio de su religion, sus costumbres, leyes y propiedades. Estos acuerdos fueron "
        "sistematicamente incumplidos: en 1502, los musulmanes del reino de Castilla fueron obligados "
        "a elegir entre la conversion al cristianismo o el exilio. La mayoria opto formalmente por la "
        "conversion, convirtiendose en moriscos. El exilio de entre 100.000 y 200.000 personas hacia "
        "el norte de Africa fue una de las grandes tragedias demograficas de la historia ibérica. "
        "Boabdil mismo marcho al exilio en Fez, donde murio en oscuridad."
    ))

    set_heading(doc, "1.4 La alianza con los Marinidas y el fin del apoyo exterior", 2)
    add_paragraph(doc, (
        "Durante el siglo XV, Granada habia intentado mantener viva la alianza con el Magreb, pero "
        "el sultanato Wattasida de Fez estaba demasiado debilitado para intervenir. El Imperio Otomano, "
        "en plena expansion (toma de Constantinopla, 1453), no pudo articular una respuesta a tiempo. "
        "Las peticiones de ayuda de los sultanes nazaries a Istanbul fueron respondidas con promesas "
        "que nunca se materializaron en apoyo militar efectivo."
    ))

    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 La economia sedera en su apogeo y crisis", 2)
    add_paragraph(doc, (
        "El siglo XV fue, paradojicamente, el periodo de mayor produccion sedera de la historia "
        "granadina. La Alpujarra concentraba decenas de miles de trabajadores de la seda, y los "
        "mercados de Granada, Almeria y Malaga exportaban tejidos que llegaban a los Paises Bajos, "
        "Italia y el Levante. Los comerciantes genoveses, especialmente la familia Centurione, "
        "controlaban el grueso del comercio exterior y financiaban incluso a los propios reyes castellanos."
    ))
    add_paragraph(doc, (
        "La guerra de 1482-1492 interrumpio brutalmente este ciclo economico: el bloqueo castellano "
        "corto las rutas comerciales, las talas de la vega impidieron las cosechas, y el asedio "
        "de ciudades claves dislocó las redes de produccion artesanal."
    ))

    set_heading(doc, "2.2 Fiscalidad e impacto de la guerra", 2)
    add_paragraph(doc, (
        "Para financiar su defensa, los sultanes nazaries del siglo XV multiplicaron las exacciones "
        "fiscales, generando malestar popular y erosionando el apoyo social a la dinastia. "
        "La presion tributaria about las comunidades campesinas de la Alpujarra fue especialmente "
        "grave. Esta crisis fiscal fue uno de los factores que debilito la cohesion interna del "
        "reino en su momento mas critico."
    ))

    set_heading(doc, "2.3 Las rutas comerciales y el rol genovés", 2)
    add_paragraph(doc, (
        "Los notarios genoveses en los puertos de Malaga y Almeria dejaron una documentacion "
        "excepcional sobre el comercio granadino del siglo XV. Sus registros muestran la variedad "
        "de productos exportados —seda, azucar, higos, pasas— y la complejidad de las redes "
        "financieras internacionales en las que Granada participaba. Cronologicamente, los mismos "
        "mercaderes que comerciaban con Granada financiaban los viajes de Colon: el mundo se "
        "expandia justo cuando Al-Andalus se contraía."
    ))

    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 La sociedad granadina en la crisis final", 2)
    add_paragraph(doc, (
        "La sociedad granadina del siglo XV presentaba una homogeneidad religiosa casi total "
        "(practicamente sin minoria cristiana ni judia), pero una gran diversidad etnica "
        "y de procedencias regionales. La presion de la guerra genero una creciente militarizacion "
        "de la vida social. Los refugiados de ciudades caidas —Ronda, Malaga, Almeria— "
        "confluyeron en Granada capital, creando una ciudad superpoblada y tensa."
    ))

    set_heading(doc, "3.2 Las comunidades mudejares en los territorios conquistados", 2)
    add_paragraph(doc, (
        "En las ciudades tomadas por los Reyes Católicos antes de 1492 —Ronda, Malaga, Almeria— "
        "se implantaron distintos regimenes para los musulmanes que permanecian: en Malaga, "
        "la resistencia llevo a la esclavizacion masiva de la poblacion; en Ronda, se aplicaron "
        "las capitulaciones. Esta variabilidad en el trato revelaba la ausencia de una politica "
        "coherente y el peso de las circunstancias militares."
    ))

    set_heading(doc, "3.3 El papel de la mujer en la sociedad nazari tardía", 2)
    add_paragraph(doc, (
        "La investigacion reciente ha prestado creciente atencion al rol de la mujer en la "
        "sociedad granadina del siglo XV. Figuras como Fatima, madre de Boabdil, que tuvo un "
        "papel politico activo en las luchas de sucesion, o las poetisas que frecuentaban "
        "la corte granadina, ilustran una presencia femenina en la vida publica que difiere "
        "de los estereotipos habituales. Los estudios de genero aplicados al periodo nazari "
        "tardio son todavia un campo academico en expansion."
    ))

    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 Los ultimos grandes cronicones nazaries", 2)
    add_paragraph(doc, (
        "Ibn al-Azraq (Granada, siglo XV) y otros cronistas de corte documentaron los ultimos "
        "decenios del reino con una urgencia narrativa palpable. Su obra, en parte perdida y en "
        "parte conservada en manuscritos del norte de Africa, es fuente primaria fundamental "
        "para la historia politica de Granada tardía. La reconstruccion de estos textos "
        "ha sido uno de los proyectos centrales de la filologia arabica espanola contemporanea."
    ))

    set_heading(doc, "4.2 La Alhambra completada: el Palacio del Partal y las ultimas obras", 2)
    add_paragraph(doc, (
        "Las obras en la Alhambra continuaron durante el siglo XV. Muhammad III habia completado ya "
        "los grandes palacios en el XIV, pero el siglo XV vio ajustes, ampliaciones y decoraciones "
        "adicionales. El conjunto alcanzo su forma definitiva antes de la conquista. Tras 1492, "
        "los Reyes Católicos respetaron inicialmente la Alhambra como palacio real propio, "
        "anadiendo luego la Capilla Real y, en el siglo XVI, el Palacio de Carlos V."
    ))

    set_heading(doc, "4.3 La poesia del fin: el genero de las elegias", 2)
    add_paragraph(doc, (
        "La literatura arabe del siglo XV en Al-Andalus esta dominada por el genero elegiaco: "
        "poemas de despedida, de nostalgia anticipada, de lamento por un mundo que se intuye "
        "condenado. Esta tradicion, que arranc desde el siglo XI con los poetas de las taifas, "
        "alcanzo en el siglo XV una intensidad especial. Algunos de estos textos son testimonios "
        "literarios de primerísimo orden sobre la experiencia del colapso civilizatorio."
    ))

    set_heading(doc, "4.4 Ciencia en el ocaso: astronomia y medicina nazaries", 2)
    add_paragraph(doc, (
        "La produccion cientifica granadina del siglo XV, aunque menos brillante que la del XIV, "
        "mantuvo la tradicion astronomica y medica heredada. Los medicos de corte siguieron "
        "compilando textos farmacologicos y botanicos. Las tablas astronomicas granadinas "
        "fueron consultadas por astrologos castellanos y portugueses, que valoraban la precisión "
        "de las observaciones andalusies. Tras la conquista, varios sabios granadinos huyeron "
        "al Magreb llevando sus libros y conocimientos, enriqueciendo las escuelas de Fez y Tunez."
    ))

    set_heading(doc, "5. Plan de Busqueda Bibliografica", 1)
    add_paragraph(doc, (
        "El siguiente plan sistematizado permite recuperar la produccion academica en WOS y Scopus "
        "para el analisis multidimensional del siglo XV nazari y la caida de Granada."
    ))
    create_search_table(doc, [
        ("1. Politico", "WOS",
         '("Al-Andalus" OR "Nasrid" OR "Granada") AND ("fall of Granada" OR "1492" OR "Boabdil" OR "capitulation*" OR "Catholic Monarchs") AND ("15th century" OR "fifteenth century") AND PY=(1980-2025)'),
        ("2. Economico", "WOS/Scopus",
         '("Nasrid" OR "Granada") AND ("silk" OR "Genoese" OR "trade" OR "fiscal" OR "economy") AND ("15th century" OR "late medieval" OR "1400-1492")'),
        ("3. Social", "WOS",
         '("Al-Andalus" OR "Nasrid Granada") AND ("Mudéjar*" OR "morisco*" OR "conversion" OR "exile" OR "expulsion" OR "demographic") AND ("15th century" OR "1492")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Nasrid" OR "Granada") AND ("Alhambra" OR "elegy" OR "poetry" OR "chronicles" OR "Ibn al-Azraq" OR "manuscript*")'),
        ("5. Transversal", "Ambas",
         'TS=("Nasrid*" OR "Al-Andalus" OR "fall of Granada" OR "Islamic Iberia" OR "Reconquista 1492") AND PY=(1980-2025)'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas", 2)
    for item in [
        "Exportar en BibTeX/RIS hacia Zotero o Mendeley. Deduplica entre WOS y Scopus.",
        "Revisar: Al-Qantara (CSIC), Journal of Medieval Iberian Studies, Arabica, Medieval Encounters.",
        "Para las Capitulaciones: consultar el Archivo General de Simancas (online) y la coleccion de Mercedes Garcia-Arenal.",
        "Para el exilio nazari en el Magreb: Archives nationales de Tunez y Biblioteca Nacional de Rabat.",
        "Cruzar resultados con historiografia de los moriscos: keyword 'Morisco*' AND 'Granada' AND '16th century'.",
        "Filtro temporal: 1985-2025. Idiomas: EN, ES, FR, AR.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo XV es el acto final de un drama de ocho siglos. La caida de Granada en 1492 "
        "no fue una derrota repentina sino el resultado de una erosion secular de territorio, "
        "recursos y cohesion politica. Y sin embargo, el reino resistio mucho mas de lo que "
        "cualquier analisis puramente geopolitico habria predicho, gracias a su solidez economica, "
        "su sofisticacion diplomatica y la fuerza de su identidad cultural."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XV ofrece la narrativa mas poderosa de toda la serie: "
        "el fin de una civilizacion y, simultaneamente, el inicio de su mito. El 'suspiro del Moro', "
        "aunque probablemente legendario, condensa en una imagen la experiencia de la perdida "
        "que resonara en la literatura, la musica y la identidad cultural de millones de personas "
        "durante los siglos siguientes."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloXV.docx")
    doc.save(out_path)
    print("Guardado: " + out_path)


# ===========================================================================
# SIGLO XVI (1500-1599)
# ===========================================================================

def create_siglo_XVI():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO XVI",
        "Los Moriscos: Herencia, Represion y Resistencia",
        "c. 1500 - 1599")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "Con la caida de Granada en 1492, Al-Andalus como Estado desaparecio, pero su civilizacion "
        "no. El siglo XVI es el siglo de los moriscos: los musulmanes convertidos forzosamente al "
        "cristianismo que permanecieron en la Peninsula manteniendo en secreto, o semisecreto, sus "
        "practicas religiosas, su lengua, su cultura y su memoria. Es un siglo de tension creciente "
        "entre la politica de asimilacion forzada de la Monarquia Hispanica y la resistencia "
        "cultural activa de comunidades que no renunciaban a su identidad."
    ))
    add_paragraph(doc, (
        "Este informe cubre el siglo XVI desde la perspectiva de la herencia viva de Al-Andalus: "
        "los moriscos como sujeto historico, su economia, su sociedad, su cultura clandestina "
        "y su relacion con el mundo otomano y norteafricano. Todas las dimensiones del periodo "
        "son analizadas con vistas a la produccion bibliometrica en WOS y Scopus para el Proyecto Noor."
    ))

    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 La conversion forzada y el problema morisco", 2)
    add_paragraph(doc, (
        "En 1502, los Reyes Católicos promulgaron la pragmatica que obligaba a los musulmanes "
        "castellanos a convertirse al cristianismo o exiliarse. La mayoria eligio la conversion "
        "formal, dando origen a la comunidad morisca. En los reinos de la Corona de Aragon, "
        "el mismo proceso se impuso entre 1525 y 1526 bajo Carlos I. Desde el punto de vista "
        "juridico, los moriscos eran cristianos nuevos; desde el punto de vista cultural y "
        "religioso, muchos seguian siendo musulmanes practicantes en privado."
    ))

    set_heading(doc, "1.2 La Gran Rebelion de las Alpujarras (1568-1571)", 2)
    add_paragraph(doc, (
        "La pragmatica de Felipe II de 1567, que prohibia el uso de la lengua arabe, los vestidos, "
        "los banos y todos los elementos culturales moriscos, desencadeno la mayor rebelion morisca "
        "del siglo: la Guerra de las Alpujarras (1568-1571). Aben Humeya (Fernando de Valor) y "
        "despues Aben Aboo lideraron una resistencia armada que se prolongó tres anos y requirio "
        "el envio de tropas veteranas de los Paises Bajos. La derrota fue seguida de la deportacion "
        "masiva de los moriscos granadinos al interior de Castilla, disolviendo su concentracion "
        "territorial para facilitar la asimilacion."
    ))

    set_heading(doc, "1.3 Felipe II y las politicas de asimilacion", 2)
    add_paragraph(doc, (
        "La politica de Felipe II hacia los moriscos oscilo entre la vigilancia inquisitorial "
        "y los intentos de asimilacion cultural programada. Se crearon escuelas para educar "
        "a los hijos de moriscos en la fe y lengua castellanas. Se promovio el matrimonio mixto. "
        "Pero el aparato inquisitorial seguia procesando a moriscos por practicas islamicas "
        "en gran numero. La tension entre el objetivo de asimilacion y el instrumento de represion "
        "fue estructural y no resuelta durante todo el siglo."
    ))

    set_heading(doc, "1.4 Las conexiones con el Imperio Otomano y el norte de Africa", 2)
    add_paragraph(doc, (
        "Los moriscos mantuvieron contactos regulares con el Imperio Otomano, que los consideraba "
        "correligionarios oprimidos y potenciales aliados en su conflicto con la Monarquia Hispanica. "
        "El corsarismo berberisco en las costas mediterraneas y atlanticas de la peninsula tenia "
        "frecuentemente informacion y colaboracion de moriscos. Esta dimension geopolitica fue "
        "real, aunque la historiografia durante siglos la exagero para justificar la expulsion. "
        "La correspondencia diplomatica otomana y las fuentes magrebies son fuentes esenciales "
        "para comprender esta red transnacional."
    ))

    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 El papel economico de los moriscos", 2)
    add_paragraph(doc, (
        "Los moriscos eran un sector economicamente indispensable en varias regiones criticas. "
        "En Valencia, constituian la practica totalidad de la mano de obra en las propiedades "
        "de la nobleza --huertas de regadio, produccion de azucar, artesanias de seda y cuero--. "
        "En Granada, controlaban el ciclo de la seda desde la cria del gusano hasta el tejido. "
        "En Aragon, eran los trabajadores especializados en la agricultura intensiva del Ebro."
    ))

    set_heading(doc, "2.2 La seda morisca: continuidad y conflicto", 2)
    add_paragraph(doc, (
        "La produccion sedera granadina continuo bajo dominio castellano, con los moriscos "
        "como productores y la Hacienda Real como principal beneficiario a traves del almojarifazgo "
        "de la seda. Esta relacion de dependencia economica era bien conocida por los funcionarios "
        "reales, que advirtieron repetidamente del coste economico que supondria la eventual "
        "expulsion de los moriscos. Sus advertencias fueron finalmente ignoradas."
    ))

    set_heading(doc, "2.3 Redes comerciales moriscas y circuitos clandestinos", 2)
    add_paragraph(doc, (
        "Los moriscos desarrollaron extensas redes comerciales que cruzaban las fronteras: "
        "mercaderes moriscos de Valencia comerciaban con Argel y Tunez, llevando y trayendo "
        "mercancias y, en ocasiones, personas. Este comercio clandestino, combinado con las "
        "remesas de los exiliados en el Magreb, creaba una economia transfronteriza que las "
        "autoridades castellanas y aragonesas nunca lograron controlar completamente."
    ))

    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 Identidad morisca: entre la asimilacion y la resistencia", 2)
    add_paragraph(doc, (
        "La identidad morisca del siglo XVI es uno de los problemas historiograficos mas debatidos "
        "de la historia ibérica moderna. Hasta que punto eran todavia musulmanes? Hasta que punto "
        "habian asimilado el catolicismo? La respuesta variaba enormemente segun la region, "
        "la generacion y el individuo. Existian moriscos plenamente asimilados que participaban "
        "normalmente en la vida cristiana; existian otros que practicaban un islam clandestino "
        "detallado y sostenido; y existia una mayoria en situaciones intermedias de identidad hibrida."
    ))

    set_heading(doc, "3.2 La Inquisicion y los moriscos", 2)
    add_paragraph(doc, (
        "El Tribunal de la Inquisicion fue el principal instrumento de presion sobre las "
        "comunidades moriscas. Sus archivos --conservados en el Archivo Historico Nacional de Madrid-- "
        "constituyen una fuente historiografica de valor extraordinario, pese a sus obvias limitaciones. "
        "Los procesos inquisitoriales contra moriscos revelan no solo practicas religiosas sino "
        "redes sociales, practicas culinarias, uso de la lengua arabe y conexiones con el exterior. "
        "La tasa de procesados moriscos fue muy superior a la de los cristianos viejos."
    ))

    set_heading(doc, "3.3 Matrimonios mixtos y frontera social", 2)
    add_paragraph(doc, (
        "La politica de matrimonios mixtos entre moriscos y cristianos viejos fue promovida "
        "oficialmente como instrumento de asimilacion. En la practica, las resistencias culturales "
        "de ambas comunidades y el estigma de la limpieza de sangre limitaron enormemente "
        "su eficacia. Los estatutos de limpieza de sangre, que privaban a los de ascendencia "
        "morisca o judia del acceso a cargos y honores, crearon una barrera estructural que "
        "impedia la asimilacion plena incluso para los moriscos dispuestos a adoptarla."
    ))

    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 La literatura aljamiada", 2)
    add_paragraph(doc, (
        "Una de las manifestaciones culturales mas fascinantes del siglo XVI morisco es la "
        "literatura aljamiada: textos en lengua castellana o aragonesa escritos en caracteres "
        "arabes. Este fenomeno, que surge como estrategia de supervivencia cultural ante la "
        "prohibicion del arabe, produjo una literatura en el cruce de dos tradiciones: "
        "relatos coranicos, vidas de profetas, textos juridicos islamicos y novelas de caballeria "
        "se escribieron todos en alfabeto arabe para una comunidad que iba perdiendo el arabe oral "
        "pero conservaba su scriptura sagrada."
    ))

    set_heading(doc, "4.2 Los manuscritos del Sacromonte: falsificacion y memoria", 2)
    add_paragraph(doc, (
        "En 1595, aparecieron en el Sacromonte de Granada unas laminas de plomo con inscripciones "
        "en arabe atribuidas a discipulos del apostol Santiago. Esta elaborada falsificacion, "
        "obra muy probablemente de moriscos cultos (Miguel de Luna, Alonso del Castillo) "
        "buscaba crear un pasado cristiano-arabe que reconciliara las dos identidades y "
        "protegiera a la comunidad morisca de la represion. El episodio, conocido como el "
        "fraude del Sacromonte, es uno de los mas complejos y fascinantes de la historia cultural ibérica."
    ))

    set_heading(doc, "4.3 La herencia arquitectonica: el arte mudéjar y la transicion al Renacimiento", 2)
    add_paragraph(doc, (
        "El siglo XVI vio la convivencia entre la herencia artistica andalusi y el nuevo lenguaje "
        "renacentista importado de Italia. El arte mudejar --sintesis de tecnicas islamicas "
        "con encargos y tipologias cristianas-- alcanzo en este siglo algunas de sus realizaciones "
        "mas complejas: la techumbre de la catedral de Teruel, los artesonados de iglesias "
        "sevillanas, los patios de las casas solariegas andaluzas. En paralelo, el Palacio de "
        "Carlos V en la Alhambra (iniciado en 1527) colocaba un edificio de planta circular "
        "estrictamente renacentista en el corazon del conjunto nazari."
    ))

    set_heading(doc, "4.4 Ciencia y herencia medica andalusi en el siglo XVI", 2)
    add_paragraph(doc, (
        "La medicina arabe tradicional, vehiculada por traducciones latinas medievales, "
        "siguio siendo texto de referencia en las universidades europeas del siglo XVI. "
        "El Canon de Avicena continuaba ensenandose en Salamanca, Padua y Montpellier. "
        "Medicos moriscos, algunos con formacion en el Magreb, ejercian en comunidades "
        "rurales espanolas. La botanica de Dioscorides, editada y comentada por el "
        "andaluz Andres Laguna (1555), sintetizaba la tradicion griega y arabe en "
        "lengua castellana, siendo uno de los textos cientificos mas influyentes del siglo."
    ))

    set_heading(doc, "5. Plan de Busqueda Bibliografica", 1)
    add_paragraph(doc, (
        "Plan sistematizado para recuperar produccion academica en WOS y Scopus sobre "
        "el legado de Al-Andalus y la comunidad morisca en el siglo XVI."
    ))
    create_search_table(doc, [
        ("1. Politico", "WOS",
         '("Morisco*" OR "Al-Andalus legacy") AND ("expulsion" OR "Inquisition" OR "Felipe II" OR "Alpujarras" OR "conversion") AND ("16th century" OR "sixteenth century") AND PY=(1980-2025)'),
        ("2. Economico", "WOS/Scopus",
         '("Morisco*" OR "Islamic Spain") AND ("silk" OR "economy" OR "trade" OR "labor" OR "Valencia" OR "Granada") AND ("16th century" OR "early modern")'),
        ("3. Social", "WOS",
         '("Morisco*") AND ("identity" OR "Inquisition" OR "limpieza de sangre" OR "assimilation" OR "social structure") AND ("16th century" OR "sixteenth century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Morisco*" OR "Al-Andalus") AND ("aljamiado" OR "Sacromonte" OR "mudejar" OR "manuscript*" OR "literature" OR "clandestine")'),
        ("5. Transversal", "Ambas",
         'TS=("Morisco*" OR "Mudéjar*" OR "Islamic Iberia" OR "converso" OR "crypto-Muslim*") AND PY=(1980-2025)'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas", 2)
    for item in [
        "Exportar en BibTeX/RIS hacia Zotero o Mendeley. Deduplica entre WOS y Scopus.",
        "Revisar: Al-Qantara (CSIC), Sharq al-Andalus, Journal of Early Modern History, Arabica.",
        "Fuente primaria clave: Archivo Historico Nacional, fondo Inquisicion (digitalizado parcialmente).",
        "Para literatura aljamiada: Biblioteca Nacional de Espana y Biblioteca de la Real Academia de la Historia.",
        "Para el Sacromonte: obras de Mercedes Garcia-Arenal y Fernando Rodriguez Mediano.",
        "Filtro: 1985-2025. Idiomas: EN, ES, FR, AR.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo XVI es el siglo de la resistencia cultural silenciosa. Los moriscos demostraron "
        "que una civilizacion no desaparece con la conquista militar: se transforma, se adapta, "
        "se camufla. La literatura aljamiada, el fraude del Sacromonte, las redes comerciales "
        "clandestinas: todas son estrategias de supervivencia identitaria de una comunidad "
        "que se niega a desaparecer."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XVI ofrece una narrativa de resiliencia extraordinaria "
        "que conecta con debates contemporaneos sobre minorías, identidad y represion cultural. "
        "Es el puente entre Al-Andalus medieval y la memoria de Al-Andalus en el mundo moderno."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloXVI.docx")
    doc.save(out_path)
    print("Guardado: " + out_path)


# ===========================================================================
# SIGLO XVII (1600-1699)
# ===========================================================================

def create_siglo_XVII():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO XVII",
        "La Expulsion de los Moriscos y el Final de la Presencia Islamica en la Peninsula",
        "c. 1600 - 1699")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "El siglo XVII comienza con la decision mas radical de la historia de la Monarquia "
        "Hispanica respecto a su herencia islamica: la expulsion total de los moriscos "
        "(1609-1614), ordenada por Felipe III y ejecutada por el Duque de Lerma. Entre "
        "300.000 y 500.000 personas --la mayor operacion de limpieza etnica-religiosa "
        "de la historia europea premodern-- fueron embarcadas forzosamente hacia el norte "
        "de Africa, Italia, Francia y el Imperio Otomano. Con esta expulsion, la presencia "
        "islamica en la Peninsula Iberica, que habia durado ocho siglos, termino formalmente."
    ))
    add_paragraph(doc, (
        "El resto del siglo XVII vio las consecuencias de esta decision: el impacto economico "
        "en las regiones que dependian de la mano de obra morisca, la formacion de comunidades "
        "de exiliados en el Magreb que conservaron durante generaciones su memoria iberica, "
        "y los primeros intentos de una historiografia que fijara el relato de lo ocurrido."
    ))

    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 La decision de la expulsion: debate y ejecucion", 2)
    add_paragraph(doc, (
        "La decision de expulsar a los moriscos no fue repentina ni unanime. Durante decadas "
        "el Consejo de Estado habia debatido entre la asimilacion y la expulsion. Los argumentos "
        "favorables a la expulsion --miedo a la quinta columna otomana, fracaso de la asimilacion, "
        "pureza religiosa-- se impusieron finalmente sobre los economicos, que senalaban el "
        "coste devastador para la agricultura valenciana y el comercio. El bando de expulsion "
        "se publico en Valencia el 22 de septiembre de 1609. Le siguieron Aragon, Castilla "
        "y Andalucia entre 1610 y 1614."
    ))

    set_heading(doc, "1.2 El impacto politico de la expulsion en la geopolitica mediterranea", 2)
    add_paragraph(doc, (
        "La expulsion tuvo efectos geopoliticos inmediatos. El norte de Africa absorbio el grueso "
        "de los exiliados: Tunez, Argelia, Marruecos. Los moriscos, que a menudo conservaban "
        "conocimientos tecnicos valiosos (agricultura de regadio, artesanias especializadas), "
        "fueron recibidos con ambivalencia: como recurso humano util y como masa empobrecida "
        "y desorientada. Algunos se integraron rapidamente; otros vivieron decadas en campamentos "
        "de refugiados. Varios grupos de exiliados granadinos en Marruecos fundaron la ciudad "
        "de Salé, desde donde dirigieron operaciones de corso contra la navegacion espanola "
        "en el Atlantico (la Republica de Bou Regreg)."
    ))

    set_heading(doc, "1.3 Las secuelas politicas internas: el debate historiografico", 2)
    add_paragraph(doc, (
        "Inmediatamente despues de la expulsion, comenzaron los debates sobre su justificacion "
        "y sus efectos. Pedro de Valencia, humanista espanol, escribio un tratado critico con "
        "la expulsion que no se publico en vida por su caracter politically sensitive. "
        "El cronista Jaime Bleda, por el contrario, proporciono la justificacion apologetica "
        "oficial. Este debate fundacional sobre la expulsion ha continuado hasta hoy y es "
        "uno de los campos mas activos de la historiografia ibérica contemporanea."
    ))

    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 El colapso economico en Valencia y otras regiones", 2)
    add_paragraph(doc, (
        "El impacto economico de la expulsion fue devastador en el Reino de Valencia, donde los "
        "moriscos representaban entre el 30% y el 35% de la poblacion total y la casi totalidad "
        "de los trabajadores agricolas. Aldeas enteras quedaron despobladas. Las senalizaciones "
        "nobiliarias que dependian de sus rentas se arruinaron. Segun calculos modernos, "
        "Valencia tardo al menos medio siglo en recuperar los niveles demograficos y economicos "
        "previos a la expulsion. El impacto fue menor, pero perceptible, en Aragon y Castilla."
    ))

    set_heading(doc, "2.2 La economia de la diaspora morisca en el Magreb", 2)
    add_paragraph(doc, (
        "Los exiliados moriscos en el norte de Africa no llegaron como mendigos: muchos traian "
        "capitales, conocimientos tecnicos y redes comerciales. En Tunez, los andaluces "
        "(como se llamaba a los exiliados moriscos en el Magreb) introdujeron o perfeccionaron "
        "tecnicas de cultivo de jasmin, naranja amarga y productos hortofruticolas. En Marruecos, "
        "las ciudades de Fez, Tetuan, Rabat y Sale recibieron flujos importantes de inmigrantes "
        "ibéricos que dejaron huella en la arquitectura, la gastronomia y las artes decorativas."
    ))

    set_heading(doc, "2.3 El corsarismo morisco como economia alternativa", 2)
    add_paragraph(doc, (
        "La Republica de Bou Regreg (Sale-Rabat), fundada en buena medida por exiliados "
        "moriscos andaluces a partir de 1614, fue durante decadas una potencia corsaria "
        "de primer orden en el Atlantico norte. Sus barcos atacaban la navegacion portuguesa, "
        "espanola, inglesa y francesa. El corso era la economia de los sin tierra, la venganza "
        "economica de quienes habian sido despojados. Miguel de Cervantes, ex cautivo en Argel, "
        "conocio de primera mano estas redes en las que ibéricos de multiple confesion convivian "
        "en una frontera cultural porosa."
    ))

    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 El trauma del exilio: testimonios y memorias", 2)
    add_paragraph(doc, (
        "La literatura de los exiliados moriscos del siglo XVII es un corpus de enorme valor "
        "historico y humano. Relatos de la travesia, lamentos por las casas y tierras perdidas, "
        "descripciones de la desorientacion en la tierra de acogida: estos textos, escritos en "
        "arabe, en aljamiado y en castellano, son testimonios directos de una de las mayores "
        "tragedias de la historia ibérica. La investigacion sobre estos materiales ha experimentado "
        "un renacimiento en las ultimas decadas."
    ))

    set_heading(doc, "3.2 La integracion de los moriscos en el Magreb y el Imperio Otomano", 2)
    add_paragraph(doc, (
        "El proceso de integracion de los exiliados en sus sociedades de acogida fue largo y "
        "complejo. En Tunez, los andaluces se concentraron en barrios propios (como el barrio "
        "andalus de Testour, con su mezquita de minarete de influencia iberica) y mantuvieron "
        "durante generaciones practicas culturales ibericas: el romance en arabe dialectal, "
        "recetas de cocina, tecnicas de bordado. En el Imperio Otomano, algunos ex moriscos "
        "alcanzaron posiciones relevantes en la administracion y el ejercito."
    ))

    set_heading(doc, "3.3 Los criptoislamicos que permanecieron: los 'Nuevos Convertidos'", 2)
    add_paragraph(doc, (
        "No todos los moriscos fueron expulsados efectivamente. Algunos lograron quedarse "
        "mediante certificados de asimilacion, matrimonios con cristianos viejos o simple "
        "evasion. Estos grupos, clandestinamente islamicos durante generaciones, son "
        "historiograficamente difíciles de rastrear pero documentados a traves de procesos "
        "inquisitoriales tardios que llegan hasta el siglo XVIII. Son el ultimo eslabon "
        "de la cadena de resistencia cultural que habia comenzado con la conversion forzada de 1502."
    ))

    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 Cervantes y la cuestion morisca", 2)
    add_paragraph(doc, (
        "Miguel de Cervantes (1547-1616) vivio y escribio en el contexto directo de la "
        "expulsion de los moriscos. El Quijote (1605-1615) contiene, en el episodio de "
        "Ricote (II, 54), uno de los testimonios literarios mas sensibles sobre la expulsion: "
        "un morisco que regresa clandestinamente a Espana a recuperar un tesoro enterrado "
        "y que, en su conversacion con Sancho, expresa la tragedia de pertenecer a dos mundos "
        "y no poder quedarse en ninguno. La vision de Cervantes sobre los moriscos ha sido "
        "objeto de intenso debate critico."
    ))

    set_heading(doc, "4.2 La historiografia del siglo XVII sobre Al-Andalus", 2)
    add_paragraph(doc, (
        "El siglo XVII vio los primeros grandes intentos historiograficos de sistematizar "
        "la historia de Al-Andalus. El Padre Higuera y otros eruditos jesuitas elaboraron "
        "falsificaciones historicas (los Falsos Cronicones) que intentaban rellenar los vacios "
        "de la historia visigoda y altomedieval con materiales inventados. Paralelamente, "
        "arabistas como Marcos Dobelio y Leon el Africano trabajaban sobre textos arabes "
        "autenticos. El Escorial contenia ya una importante coleccion de manuscritos arabes "
        "capturados o adquiridos, base de futura investigacion arabista."
    ))

    set_heading(doc, "4.3 La herencia andalusi en el Barroco espanol", 2)
    add_paragraph(doc, (
        "La herencia estetica de Al-Andalus pervivio en el arte barroco espanol de "
        "multiples formas: en la decoracion de azulejos, en el uso del arco mixtilineo, "
        "en los patios con fuentes centrales de las casas senoriales andaluzas, en "
        "la musica de raiz moruna que influyo en el flamenco inicial. Esta presencia "
        "estetica, no siempre consciente ni reconocida, esta siendo sistematicamente "
        "documentada por la historia del arte contemporanea."
    ))

    set_heading(doc, "4.4 La memoria andalusi en la cultura magrebí del siglo XVII", 2)
    add_paragraph(doc, (
        "En el Magreb, los exiliados andaluces del siglo XVII eran portadores de una "
        "cultura de alta sofisticacion: musica clasica andaluza (nuba), arquitectura "
        "domestica refinada, gastronomia elaborada, artesania textil y del cuero. "
        "Estos aportes se integraron profundamente en las culturas de Tunez, Marruecos "
        "y Argelia, donde la influencia andaluza sigue siendo perceptible hoy. "
        "La musica andaluza tradicional del Magreb --uno de los patrimonios culturales "
        "inmateriales mas preciosos del Mediterraneo-- tiene su origen directo en "
        "esta transmision del siglo XVII."
    ))

    set_heading(doc, "5. Plan de Busqueda Bibliografica", 1)
    add_paragraph(doc, (
        "Plan sistematizado para recuperar produccion academica en WOS y Scopus sobre "
        "la expulsion de los moriscos, la diaspora andalusi y la herencia de Al-Andalus en el siglo XVII."
    ))
    create_search_table(doc, [
        ("1. Politico", "WOS",
         '("Morisco*" OR "Al-Andalus") AND ("expulsion" OR "1609" OR "Felipe III" OR "diaspora" OR "exile") AND ("17th century" OR "seventeenth century") AND PY=(1980-2025)'),
        ("2. Economico", "WOS/Scopus",
         '("Morisco*" OR "Andalusian exile") AND ("economic impact" OR "agriculture" OR "Valencia" OR "corsair*" OR "Bou Regreg" OR "Maghreb") AND ("17th century" OR "1609-1614")'),
        ("3. Social", "WOS",
         '("Morisco*" OR "Andalusian diaspora") AND ("integration" OR "memory" OR "identity" OR "North Africa" OR "Ottoman" OR "Tunis") AND ("17th century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Morisco*" OR "Al-Andalus legacy") AND ("Cervantes" OR "aljamiado" OR "Andalusian music" OR "nuba" OR "baroque" OR "heritage")'),
        ("5. Transversal", "Ambas",
         'TS=("Morisco*" OR "Andalusian exile" OR "Islamic Iberia legacy" OR "crypto-Muslim*" OR "Maghreb Andalusian") AND PY=(1980-2025)'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas", 2)
    for item in [
        "Exportar en BibTeX/RIS hacia Zotero o Mendeley. Deduplica entre WOS y Scopus.",
        "Revisar: Sharq al-Andalus, Al-Qantara, Journal of Early Modern History, Mediterranean Historical Review.",
        "Para el corsarismo: obras de Bartolome y Lucile Bennassar; para la economia valenciana: Henri Lapeyre.",
        "Fuentes en el norte de Africa: Archives Nationales de Tunez; Direction des Archives Royales (Rabat).",
        "Para la herencia musical: UNESCO Intangible Heritage database; trabajos de Dwight Reynolds sobre musica andaluza.",
        "Filtro: 1985-2025. Idiomas: EN, ES, FR, AR.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo XVII es el siglo de la ruptura definitiva y, paradojicamente, de la "
        "supervivencia cultural mas tenaz. La expulsion de los moriscos no liquido "
        "la herencia de Al-Andalus: la disperso por el Mediterraneo, donde siguio "
        "germinando en formas distintas. La musica andaluza del Magreb, la arquitectura "
        "de Tetuan, la gastronomia de Tunez y los apellidos ibericos de familias arabas "
        "son testigos vivos de que una civilizacion puede perder su Estado y conservar "
        "su alma."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XVII ofrece la historia de la diaspora: "
        "el Al-Andalus disperso, fragmentado y transformado que continua existiendo "
        "mas alla de sus fronteras originales. Es quizas el capitulo mas universal "
        "del relato andalusi, el que mas directamente conecta con experiencias de "
        "refugio, exilio e identidad propias del mundo contemporaneo."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloXVII.docx")
    doc.save(out_path)
    print("Guardado: " + out_path)


# ===========================================================================
# SIGLO XVIII (1700-1799)
# ===========================================================================

def create_siglo_XVIII():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO XVIII",
        "La Memoria Ilustrada: Orientalismo Temprano y Redescubrimiento Academico",
        "c. 1700 - 1799")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "En el siglo XVIII, Al-Andalus ya no existe como entidad politica: la expulsion de los "
        "moriscos en 1609-1614 habia eliminado la ultima comunidad que mantenia una conexion "
        "viviente con la civilizacion iberica islamica. Y sin embargo, el siglo XVIII es el siglo "
        "en que Al-Andalus comienza a ser redescubierta como objeto de fascinacion intelectual "
        "y estetica. La Ilustracion europea, con su interes por la historia comparada de las "
        "civilizaciones, y los viajeros britanicos, franceses y alemanes que recorrieron la "
        "Peninsula, pusieron los cimientos del orientalismo andalusi que floreceria en el siglo XIX."
    ))
    add_paragraph(doc, (
        "Este informe analiza el siglo XVIII desde la perspectiva de la herencia de Al-Andalus: "
        "la historiografia ilustrada, la situacion de las comunidades de origen andalusi en el "
        "Magreb, el inicio de la arqueologia arabica, y el impacto estetico del 'estilo morisco' "
        "en las artes decorativas europeas."
    ))

    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 La politica borbonica y el legado arabe en Espana", 2)
    add_paragraph(doc, (
        "La dinastia borbonica, instalada en el trono espanol tras la Guerra de Sucesion "
        "(1701-1714), impulso una modernizacion ilustrada que incluia una nueva atencion "
        "a la historia nacional. Los Borbones, especialmente Carlos III, fomentaron "
        "la creacion de academias e instituciones de investigacion que comenzaron a "
        "incorporar el estudio del pasado islamico peninsular como parte de la historia "
        "de Espana. La Real Academia de la Historia inicio la publicacion de fuentes "
        "arabes en traduccion castellana."
    ))

    set_heading(doc, "1.2 Las comunidades andalusies en el Magreb: situacion politica en el siglo XVIII", 2)
    add_paragraph(doc, (
        "Las comunidades de origen andalusi en el norte de Africa llevaban un siglo de "
        "integracion parcial en sus sociedades de acogida. En Tunez, los andaluces "
        "habian alcanzado posiciones prominentes en la artesania, el comercio y, "
        "en algunos casos, la administracion. En Marruecos, el sultan Mulay Ismail "
        "(r. 1672-1727) habia utilizado a descendientes de andalusies como tropas "
        "de elite y tecnicos. A lo largo del siglo XVIII, la distincion entre "
        "andalusies y poblaciones locales fue diluyendose, aunque la memoria iberica "
        "se conservo en cofradias, barrios y tradiciones especificas."
    ))

    set_heading(doc, "1.3 Las relaciones diplomaticas hispano-magrebies y la memoria de Al-Andalus", 2)
    add_paragraph(doc, (
        "Las negociaciones diplomaticas entre Espana y los sultanatos del norte de Africa "
        "estuvieron frecuentemente cargadas de referencias a la historia andalusi. "
        "Los sultanes marroquies invocaban la herencia de Al-Andalus para reclamar "
        "enclaves como Ceuta y Melilla. Las negociaciones de rescate de cautivos, "
        "que continuaron durante el siglo XVIII, mantuvieron vivos los contactos "
        "humanos entre las dos orillas del Estrecho."
    ))

    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 La economia de las ciudades andalusies del Magreb", 2)
    add_paragraph(doc, (
        "Las ciudades de fuerte impronta andaluza en el Magreb --Tetuan, Chefchaouen, "
        "Testour, Soliman-- mantuvieron en el siglo XVIII una economia artesanal "
        "caracteristica heredada de la tradicion iberica: tejidos de seda y lana de "
        "alta calidad, cuero trabajado (la cordoban leather europea toma su nombre "
        "de aqui), ceramica, metalurgia del cobre. Estos productos circulaban en "
        "redes comerciales mediterraneas en las que europeos, judios sefardies y "
        "mercaderes arabes se encontraban."
    ))

    set_heading(doc, "2.2 El comercio mediterraneo y la herencia de las rutas andalusies", 2)
    add_paragraph(doc, (
        "Las rutas comerciales del Mediterraneo occidental del siglo XVIII reflejaban "
        "aun la herencia de los flujos andalusies medievales. Los puertos de Argel, "
        "Tunez y Tripoli eran escalas esenciales en la navegacion entre el Atlantico "
        "y el Levante. Los productos de origen o tecnica andalusi --tejidos, ceramica, "
        "especias reexportadas-- seguian circulando, aunque bajo etiquetas genovesas, "
        "marsellesas o britanicas."
    ))

    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 Las cofradias andalusies en el Magreb: preservacion de la identidad", 2)
    add_paragraph(doc, (
        "Las comunidades de origen andalusi en el norte de Africa del siglo XVIII se organizaban "
        "frecuentemente en torno a cofradias religiosas y asociaciones de compatriatas. "
        "Estas instituciones cumplian funciones de solidaridad mutual, preservacion cultural "
        "y mantenimiento de la memoria colectiva. La investigacion etnografica e historica "
        "sobre estas cofradias en el siglo XX ha revelado la extraordinaria persistencia "
        "de practicas culturales ibericas en el Magreb varios siglos despues de la expulsion."
    ))

    set_heading(doc, "3.2 Los viajeros europeos y la imagen social de Al-Andalus", 2)
    add_paragraph(doc, (
        "Los miles de viajeros europeos que visitaron Espana en el siglo XVIII constituyen "
        "una fuente sociologica de primer orden. Sus relatos, publicados en Londres, Paris "
        "y Berlín, construyeron una imagen de Al-Andalus como mundo perdido de refinamiento "
        "y exotismo. Henry Swinburne, Richard Twiss y otros visitaron la Alhambra "
        "cuando era literalmente una ruina habitada por gitanos y vagabundos, y "
        "sus descripciones contribuyeron a despertar el interes europeo por la "
        "arquitectura islamica."
    ))

    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 La Ilustracion y los arabistas espanoles", 2)
    add_paragraph(doc, (
        "El siglo XVIII vio surgir la primera generacion de arabistas espanoles en sentido moderno. "
        "Miguel Casiri (1710-1791), sacerdote maronita al servicio de la corona espanola, "
        "catalogo los manuscritos arabes del Escorial y publico la Bibliotheca Arabico-Hispana "
        "Escurialensis (1760-1770), primera gran obra de arabismo espanol sistematico. "
        "Juan de Iriarte y otros ilustrados promovieron el estudio de la historia arabe "
        "peninsular como parte de la historia de Espana."
    ))

    set_heading(doc, "4.2 El orientalismo literario europeo y Al-Andalus", 2)
    add_paragraph(doc, (
        "La moda orientalista del siglo XVIII en Europa --estimulada por la traduccion de "
        "Las mil y una noches al frances por Galland (1704-1717) y los proyectos de "
        "Montesquieu, Voltaire y otros-- encontro en Al-Andalus un objeto de fascinacion "
        "particular. La imagen del moro civilizado y cultivado, en contraste con el "
        "fanatismo de la Inquisicion, sirvio a los ilustrados como herramienta critica "
        "contra el clericalismo europeo. Esta utilizacion ideologica de Al-Andalus "
        "es un capitulo fascinante de la historia intelectual de la Ilustracion."
    ))

    set_heading(doc, "4.3 La musica andaluza del Magreb: codificacion y transmision", 2)
    add_paragraph(doc, (
        "El siglo XVIII fue un periodo crucial para la codificacion de la musica andaluza "
        "tradicional del Magreb. Las nubas --suites musicales de origen andalusi-- fueron "
        "recopiladas y sistematizadas en Marruecos, Argelia y Tunez, donde habian evolucionado "
        "con influencias locales. Esta codificacion fue esencial para la supervivencia del "
        "repertorio hasta el siglo XX. Las investigaciones musicologicas modernas han podido "
        "identificar elementos andalusies originales en estas tradiciones magrebies."
    ))

    set_heading(doc, "4.4 La arquitectura neomorisca: inicio de una nueva estetica", 2)
    add_paragraph(doc, (
        "El redescubrimiento de la arquitectura islamica andalusi genero en el siglo XVIII "
        "los primeros experimentos de revival neomoriscos en Europa. En los jardines paisajistas "
        "ingleses aparecieron templetes de estilo morisco. En Francia y Alemania, la moda del "
        "turquerie y el chinoiserie incluia elementos andalusies. Espana sera en el siglo XIX "
        "el centro del revival neomoriscos mas intenso, pero sus raices estan en este "
        "siglo XVIII de curiosidad ilustrada."
    ))

    set_heading(doc, "5. Plan de Busqueda Bibliografica", 1)
    add_paragraph(doc, (
        "Plan sistematizado en WOS y Scopus para el analisis de la herencia de Al-Andalus "
        "y su recepcion intelectual y cultural en el siglo XVIII."
    ))
    create_search_table(doc, [
        ("1. Politico", "WOS",
         '("Al-Andalus" OR "Morisco*" OR "Andalusian diaspora") AND ("Maghreb" OR "North Africa" OR "memory" OR "Bourbon" OR "18th century") AND PY=(1980-2025)'),
        ("2. Economico", "WOS/Scopus",
         '("Andalusian" OR "Morisco*") AND ("economy" OR "artisan*" OR "trade" OR "Mediterranean" OR "Maghreb") AND ("18th century" OR "early modern period")'),
        ("3. Social", "WOS",
         '("Andalusian diaspora" OR "Morisco*" OR "Al-Andalus") AND ("identity" OR "memory" OR "cofrades" OR "North Africa" OR "community") AND ("18th century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Islamic Iberia") AND ("Orientalism" OR "Casiri" OR "Alhambra" OR "arabist*" OR "Enlightenment" OR "Andalusian music" OR "nuba")'),
        ("5. Transversal", "Ambas",
         'TS=("Al-Andalus legacy" OR "Andalusian heritage" OR "Orientalism" OR "Islamic Iberia" OR "Morisco diaspora") AND PY=(1980-2025)'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas", 2)
    for item in [
        "Exportar en BibTeX/RIS hacia Zotero o Mendeley. Deduplica entre WOS y Scopus.",
        "Revisar: Al-Qantara, Journal of the History of Ideas, Eighteenth-Century Studies, IJMES.",
        "Para Casiri y el arabismo ilustrado: Biblioteca Nacional de Espana y Real Academia de la Historia.",
        "Para viajeros europeos: Bodleian Library (Oxford) y BnF (Paris) tienen colecciones de libros de viajes.",
        "Para musica andaluza: Revue de Musicologie y Yearbook for Traditional Music.",
        "Filtro: 1985-2025. Idiomas: EN, ES, FR, AR, DE.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo XVIII es el siglo en que Al-Andalus pasa de ser una presencia viva, "
        "aunque disminuida, a convertirse en un objeto de estudio y una fuente de "
        "fascinacion estetica. Los ilustrados europeos encontraron en la civilizacion "
        "andalusi un espejo en el que reflexionar sobre los valores del conocimiento, "
        "la tolerancia y la convivencia --aunque frecuentemente proyectando sus propios "
        "deseos en un pasado idealizado."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XVIII es el capitulo del redescubrimiento: "
        "como Europa descubrio o reinvento a Al-Andalus, y como esa reinvencion "
        "contribuyo a crear el mito romantico que el siglo XIX llevaria a su apogeo."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloXVIII.docx")
    doc.save(out_path)
    print("Guardado: " + out_path)


# ===========================================================================
# SIGLO XIX (1800-1899)
# ===========================================================================

def create_siglo_XIX():
    doc = Document()
    set_margins(doc)
    build_cover(doc,
        "AL-ANDALUS EN EL SIGLO XIX",
        "El Mito Romantico: Orientalismo, Arqueologia y la Reinvencion de Al-Andalus",
        "c. 1800 - 1899")

    set_heading(doc, "Introduccion", 1)
    add_paragraph(doc, (
        "El siglo XIX es el siglo de la invencion del mito romantico de Al-Andalus. "
        "A traves de la literatura, la pintura, la musica y la arquitectura, el mundo europeo "
        "y latinoamericano construyo una imagen de la civilizacion andalusi como paraiso perdido "
        "de belleza, sensualidad, tolerancia y sabiduria. Washington Irving escribio sus "
        "Cuentos de la Alhambra desde el interior del palacio nazari; Owen Jones documento "
        "sus azulejos y mocárabes en una obra de influencia enorme sobre el diseno moderno; "
        "la opera y la zarzuela dieron al publico europeo visiones romanticas del amor morisco. "
        "Este proceso de mitificacion tuvo consecuencias hondas: a la vez que fijaba "
        "una imagen distorsionada, despertaba un interes cientifico real que fundaria "
        "la arabistica moderna."
    ))
    add_paragraph(doc, (
        "Este informe analiza el siglo XIX desde la doble perspectiva de la mitificacion romantica "
        "y del inicio de la investigacion cientifica sobre Al-Andalus, con especial atencion "
        "a las implicaciones para el Proyecto Noor y el perfil bibliometrico en WOS/Scopus."
    ))

    set_heading(doc, "1. Ambito Politico", 1)

    set_heading(doc, "1.1 La cuestion andalusi en el nacionalismo espanol decimononico", 2)
    add_paragraph(doc, (
        "La construccion del nacionalismo espanol en el siglo XIX tuvo que enfrentar la "
        "cuestion de como integrar ocho siglos de historia islamica en el relato identitario "
        "nacional. Las respuestas fueron diversas y reveladoras: algunos historiadores, como "
        "Francisco Javier Simonet, optaron por una lectura de la historia andalusi como "
        "ocupacion extranjera que habia que superar; otros, como Francisco Codera y Zaidin, "
        "fundador de la arabistica cientifica espanola, abogaban por el reconocimiento de "
        "Al-Andalus como parte esencial de la historia nacional."
    ))

    set_heading(doc, "1.2 El colonialismo espanol en Marruecos y la memoria andalusi", 2)
    add_paragraph(doc, (
        "La expansion colonial espanola en el norte de Africa a partir de la Guerra de "
        "Africa (1859-1860) activo la memoria de Al-Andalus de formas complejas. "
        "Espana se presentaba como heredera de una relacion historica con el Magreb; "
        "los intelectuales espanoles descubrieron en las ciudades marroquies huellas "
        "de la arquitectura y la cultura andaluza. Esta dimension colonial del orientalismo "
        "espanol es uno de los campos mas activos de la historiografia reciente."
    ))

    set_heading(doc, "1.3 Al-Andalus en el imaginario de las independencias latinoamericanas", 2)
    add_paragraph(doc, (
        "El romanticismo andalusista cruzo el Atlantico y tuvo un impacto significativo "
        "en la literatura y la iconografia de las republicas latinoamericanas del siglo XIX. "
        "Poetas como Jose Zorrilla, cuya obra Don Juan Tenorio funde el mito del seductor "
        "con el ambiente morisco sevillano, fueron enormemente populares en America Latina. "
        "La arquitectura neomorisca se difundio en edificios publicos desde Mexico hasta Argentina."
    ))

    set_heading(doc, "2. Ambito Economico", 1)

    set_heading(doc, "2.1 La industria del turismo y la Alhambra como destino", 2)
    add_paragraph(doc, (
        "El siglo XIX vio el nacimiento del turismo moderno en Espana, con la Alhambra "
        "como principal destino. Viajeros britanicos, franceses y alemanes inundaron Granada "
        "en busca de la experiencia romantica descrita en la literatura y la pintura. "
        "Esta afluencia turistica genero una economia local alrededor del monumento: "
        "guias, posadas, artesanos que producian replicas. El turismo orientalista del "
        "siglo XIX sento las bases de la industria turistica andaluza contemporanea."
    ))

    set_heading(doc, "2.2 La economia de las comunidades andalusies en el Magreb en el siglo XIX", 2)
    add_paragraph(doc, (
        "Las comunidades de origen andalusi en el Magreb del siglo XIX eran ya "
        "completamente integradas en sus sociedades de acogida, aunque mantuvieran "
        "rasgos culturales diferenciados. En el contexto del colonialismo europeo "
        "en Africa del Norte --Francia en Argelia desde 1830, Espana en el norte "
        "de Marruecos— estas comunidades navegaron entre identidades complejas: "
        "su herencia andaluza, su ciudadania marroquí o tunecina, y la presion "
        "colonial europea."
    ))

    set_heading(doc, "3. Ambito Social", 1)

    set_heading(doc, "3.1 El orientalismo como construccion social del otro", 2)
    add_paragraph(doc, (
        "La tesis de Edward Said (Orientalismo, 1978) sobre la construccion occidental "
        "del mundo islamico como alteridad exotica y degradada tiene en el caso andalusi "
        "una expresion especialmente compleja. El orientalismo europeo del siglo XIX "
        "sobre Al-Andalus era a la vez fascinado y condescendiente: admiraba la civilizacion "
        "perdida y justificaba la superioridad europea presente. Esta tension ha sido "
        "analizada en profundidad por la historia cultural y los estudios poscoloniales "
        "de las ultimas decadas."
    ))

    set_heading(doc, "3.2 Las comunidades gitanas y la invencion del flamenco", 2)
    add_paragraph(doc, (
        "El siglo XIX vio la cristalizacion del flamenco como forma musical y coreografica "
        "reconocible y codificada. Aunque sus origenes son complejos y debatidos, "
        "la musicologia actual reconoce aportes andalusies (melodias modales, ornamentacion "
        "vocal), gitanos y judios sefardíes en su formacion. La vision romantica del "
        "flamenco como expresion directa del alma morisca fue una de las construcciones "
        "del siglo XIX que, aunque historiograficamente simplificadora, contribuyo "
        "a fijar la identidad del genero."
    ))

    set_heading(doc, "4. Ambito Cultural y Cientifico", 1)

    set_heading(doc, "4.1 Washington Irving y los Cuentos de la Alhambra (1832)", 2)
    add_paragraph(doc, (
        "Washington Irving (1783-1859), diplomatico y escritor estadounidense, vivio "
        "durante meses en la Alhambra y publico en 1832 los Cuentos de la Alhambra, "
        "obra que tuvo un impacto cultural de primera magnitud en la imagen de Granada "
        "y Al-Andalus en el mundo anglosajono. Sus relatos, mezcla de historia, leyenda "
        "y invencion, convirtieron la Alhambra en un icono de la cultura romantica "
        "internacional. La obra sigue siendo uno de los libros mas vendidos sobre Granada."
    ))

    set_heading(doc, "4.2 Owen Jones y la gramatica del ornamento (1856)", 2)
    add_paragraph(doc, (
        "El arquitecto y disenador britanico Owen Jones (1809-1874) visito la Alhambra "
        "en 1833 y 1837, realizando mediciones y calcos precisos de su decoracion. "
        "Su publicacion Plans, Elevations, Sections and Details of the Alhambra (1842) "
        "y, especialmente, The Grammar of Ornament (1856) --que dedicaba cuatro laminas "
        "al ornamento morisco-- tuvieron una influencia decisiva en el diseno industrial, "
        "la arquitectura y las artes decorativas victorianas y posteriormente en el "
        "Art Nouveau y el modernismo."
    ))

    set_heading(doc, "4.3 El arabismo cientifico espanol: Codera, Ribera y la Escuela arabista", 2)
    add_paragraph(doc, (
        "La segunda mitad del siglo XIX vio la fundacion de la arabistica cientifica espanola "
        "de la mano de Francisco Codera y Zaidin (1836-1917) y sus discipulos Julian Ribera "
        "y Tirador y Miguel Asin Palacios. Formados en filologia clasica y lingüistica comparada, "
        "estos arabistas aplicaron metodos rigurosos al estudio de las fuentes arabes, "
        "publicaron ediciones criticas de textos medievales y establecieron las bases "
        "metodologicas que aun sostienen la arabistica hispana. Su obra es el hito "
        "fundacional de la disciplina."
    ))

    set_heading(doc, "4.4 La arquitectura neomorisca: del revival romantico al eclecticismo", 2)
    add_paragraph(doc, (
        "El siglo XIX vio la explosion del estilo neomoriscos en la arquitectura y las "
        "artes decorativas de Europa y America. En Espana, la Sala de los Embajadores "
        "del Alcazar de Sevilla y edificios como el Casino de la Sociedad Literaria de "
        "Granada fueron construidos o renovados en estilo neomoriscos. En Estados Unidos, "
        "la moda morisca influyo en edificios teatrales y de entretenimiento. "
        "En Alemania, Gottfried Semper y otros arquitectos estudiaron la Alhambra "
        "como modelo de sintesis entre estructura y decoracion."
    ))

    set_heading(doc, "5. Plan de Busqueda Bibliografica", 1)
    add_paragraph(doc, (
        "Plan sistematizado en WOS y Scopus para la produccion academica sobre el "
        "orientalismo andalusi, el romanticsimo y el arabismo cientifico del siglo XIX."
    ))
    create_search_table(doc, [
        ("1. Politico", "WOS",
         '("Al-Andalus" OR "Islamic Iberia") AND ("nationalism" OR "Orientalism" OR "colonialism" OR "Morocco" OR "Spain 19th century") AND PY=(1980-2025)'),
        ("2. Economico", "WOS/Scopus",
         '("Alhambra" OR "Granada" OR "Al-Andalus") AND ("tourism" OR "heritage" OR "economy" OR "Maghreb" OR "colonialism") AND ("19th century")'),
        ("3. Social", "WOS",
         '("Al-Andalus" OR "Orientalism" OR "Morisco*") AND ("Said" OR "identity" OR "flamenco" OR "social construction" OR "Andalusia") AND ("19th century")'),
        ("4. Cultural", "WOS/Scopus",
         '("Al-Andalus" OR "Alhambra" OR "Moorish revival") AND ("Washington Irving" OR "Owen Jones" OR "arabism" OR "Codera" OR "neo-Moorish" OR "romanticism")'),
        ("5. Transversal", "Ambas",
         'TS=("Al-Andalus legacy" OR "Moorish revival" OR "Orientalism" OR "Islamic Iberia heritage" OR "Andalusian heritage" OR "arabist*") AND PY=(1980-2025)'),
    ])

    set_heading(doc, "5.1 Recomendaciones metodologicas", 2)
    for item in [
        "Exportar en BibTeX/RIS hacia Zotero o Mendeley. Deduplica entre WOS y Scopus.",
        "Revisar: Al-Qantara, Journal of Spanish Cultural Studies, Victorian Studies, IJMES, Muqarnas.",
        "Para Washington Irving: The Complete Works of Washington Irving (ed. Wisconsin U.P.).",
        "Para Owen Jones: Victoria & Albert Museum Archive y el National Art Library (Londres).",
        "Para arabismo espanol: Revista del Instituto Egipcio de Estudios Islamicos en Madrid.",
        "Para orientalismo y poscolonialismo: cruzar con palabras clave 'Said' AND 'Andalus' en bases de datos de humanidades.",
        "Filtro: 1985-2025. Idiomas: EN, ES, FR, AR, DE.",
    ]:
        add_bullet(doc, item)

    set_heading(doc, "Conclusion", 1)
    add_paragraph(doc, (
        "El siglo XIX representa la paradoja final del ciclo de Al-Andalus: cuanto mas se "
        "alejaba en el tiempo la civilizacion original, mas intensamente era reinventada "
        "y celebrada. El mito romantico de Al-Andalus --con toda su idealizacion y sus "
        "distorsiones-- tuvo consecuencias reales y perdurables: provoco el primer impulso "
        "de conservacion de la Alhambra (antes de Irving habia sido abandonada a su suerte), "
        "fundo la arabistica cientifica y genero una conciencia de Al-Andalus como "
        "patrimonio comun que trasciende las fronteras nacionales y religiosas."
    ))
    add_paragraph(doc, (
        "Para el Proyecto Noor, el siglo XIX ofrece la historia de como un legado se "
        "convierte en mito y como ese mito, a su vez, impulsa el conocimiento cientifico. "
        "Es el capitulo que conecta la historia medieval de Al-Andalus con las preguntas "
        "que el mundo contemporaneo sigue haciendose sobre convivencia, tolerancia e identidad "
        "multicultural."
    ))

    out_path = os.path.join(OUTPUT_DIR, "Plan de Investigacion AlAndalus_SigloXIX.docx")
    doc.save(out_path)
    print("Guardado: " + out_path)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    print("Generando Planes de Investigacion Al-Andalus (Siglos XV-XIX) - Proyecto Noor")
    print("=" * 70)
    create_siglo_XV()
    create_siglo_XVI()
    create_siglo_XVII()
    create_siglo_XVIII()
    create_siglo_XIX()
    print("=" * 70)
    print("COMPLETADO. Documentos guardados en:")
    print(OUTPUT_DIR)
